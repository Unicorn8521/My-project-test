import datetime
import json
import os
import sys
import time
import tkinter as tk
import uuid
from tkinter import font as tkfont
from tkinter import ttk, filedialog, messagebox, simpledialog, scrolledtext

import keyboard
from PIL import Image, ImageTk, ImageGrab
from docx import Document
from docx.shared import Inches, Pt  # 新增Pt导入
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 新增对齐方式导入
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image as RLImage, Spacer

# 确保中文显示正常
try:
    import matplotlib

    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]
except ImportError:
    pass


class DocxReportGenerator:
    @staticmethod
    def format_duration(seconds):
        minutes, seconds = divmod(seconds, 60)
        hours, minutes = divmod(minutes, 60)
        if hours > 0:
            return f"{hours}h{minutes}m"
        elif minutes > 0:
            return f"{minutes}m{seconds}s"
        else:
            return f"{seconds}s"

    @staticmethod
    def generate_docx_report(session, save_path):
        doc = Document()
        # 报告主标题（保留标题格式）
        doc.add_heading("操作记录报告", 0)

        # 基本信息部分（保留标题格式）
        doc.add_heading("一、基本信息", level=1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.cell(0, 0).text = "操作名称"
        info_table.cell(0, 1).text = session["name"]
        info_table.cell(1, 0).text = "操作描述"
        info_table.cell(1, 1).text = session["description"]
        info_table.cell(2, 0).text = "开始时间"
        info_table.cell(2, 1).text = session["start_time"]
        info_table.cell(3, 0).text = "结束时间"
        info_table.cell(3, 1).text = session["end_time"]
        info_table.cell(4, 0).text = "操作时长"
        info_table.cell(4, 1).text = DocxReportGenerator.format_duration(session["duration"])

        # 操作步骤总标题（保留标题格式）
        doc.add_heading("二、操作步骤", level=1)
        doc.add_paragraph(f"共 {len(session['captures'])} 步操作")
        doc.add_paragraph("")  # 增加空行分隔

        for i, capture in enumerate(session["captures"], 1):
            # 步骤标题使用普通段落+加粗样式（不使用标题格式）
            step_title = doc.add_paragraph()
            step_title_run = step_title.add_run(f"步骤 {i}")
            step_title_run.bold = True  # 加粗
            step_title_run.font.size = Pt(14)  # 设置字号
            step_title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐

            # 步骤描述
            doc.add_paragraph(capture["description"])

            try:
                if os.path.exists(capture["image_path"]):
                    # 添加图片
                    doc.add_picture(capture["image_path"], width=Inches(6))
                    # 图片标题（优化字体大小单位）
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
                    run = caption.add_run(f"图 {i}: {capture['description']}")
                    run.font.size = Pt(11)  # 使用Pt单位更合适
                    run.italic = False
                else:
                    doc.add_paragraph(f"[图片文件不存在: {capture['image_path']}]")
            except Exception as e:
                doc.add_paragraph(f"[加载图片失败: {str(e)}]")

            # 步骤之间添加分页符
            doc.add_page_break()

        doc.save(save_path)
        return True


class ScreenCaptureTool:
    def __init__(self, root):
        self.root = root
        self.root.title("屏幕捕捉工具 - 就绪")
        self.root.geometry("600x400")
        self.root.resizable(True, True)

        # 数据存储配置
        self.current_session_id = None
        self.current_session = {
            "name": "",
            "description": "",
            "start_time": "",
            "end_time": "",
            "duration": 0,
            "captures": []
        }
        self.history_sessions = []
        self.sessions_dir = os.path.join(os.path.expanduser("~"), ".screen_capture_sessions")
        os.makedirs(self.sessions_dir, exist_ok=True)

        # 状态变量
        self.is_capturing = False
        self.hotkey = "ctrl+alt+o"
        self.hotkey_obj = None  # 存储热键对象（修复热键移除问题）
        self.default_font = tkfont.nametofont("TkDefaultFont")
        self.default_font.configure(family=["SimHei", "WenQuanYi Micro Hei", "Heiti TC"])
        self.root.option_add("*Font", self.default_font)

        # 加载历史记录
        self.load_history_sessions()

        # 绑定关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_main_window_close)

        # 创建界面
        self.create_main_interface()

    def create_main_interface(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题与状态
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=5)
        ttk.Label(title_frame, text="屏幕捕捉工具", font=("SimHei", 16)).pack(side=tk.LEFT)
        self.status_label = ttk.Label(title_frame, text="状态：就绪", font=("SimHei", 10), foreground="#666")
        self.status_label.pack(side=tk.RIGHT)

        # 控制按钮
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(pady=10)
        self.start_btn = ttk.Button(control_frame, text="开始捕捉", command=self.start_capture)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        self.stop_btn = ttk.Button(control_frame, text="停止捕捉", command=self.stop_capture, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=5)

        # 历史记录区域
        ttk.Label(main_frame, text="历史记录", font=("SimHei", 12)).pack(anchor=tk.W, pady=5)
        history_frame = ttk.Frame(main_frame)
        history_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 历史记录表格
        columns = ("name", "date", "duration", "captures")
        self.history_tree = ttk.Treeview(history_frame, columns=columns, show="headings")
        self.history_tree.heading("name", text="操作名称")
        self.history_tree.heading("date", text="开始时间")
        self.history_tree.heading("duration", text="时长")
        self.history_tree.heading("captures", text="截图数")
        self.history_tree.column("name", width=180)
        self.history_tree.column("date", width=180)
        self.history_tree.column("duration", width=80, anchor=tk.CENTER)
        self.history_tree.column("captures", width=80, anchor=tk.CENTER)
        self.history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 滚动条
        scrollbar = ttk.Scrollbar(history_frame, orient=tk.VERTICAL, command=self.history_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.history_tree.configure(yscrollcommand=scrollbar.set)

        # 历史记录右键菜单
        self.history_menu = tk.Menu(self.root, tearoff=0)
        self.history_menu.add_command(label="重命名", command=self.rename_session)
        self.history_menu.add_command(label="删除", command=self.delete_session)

        # 绑定事件
        self.history_tree.bind("<Double-1>", self.open_session)
        self.history_tree.bind("<Button-3>", self.show_history_menu)

        # 更新历史记录列表
        self.update_history_list()

    def show_history_menu(self, event):
        item = self.history_tree.identify_row(event.y)
        if item:
            self.history_tree.selection_set(item)
            self.history_menu.post(event.x_root, event.y_root)

    def rename_session(self):
        selected_item = self.history_tree.selection()
        if not selected_item:
            messagebox.showwarning("提示", "请先选择一条历史记录")
            return
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                new_name = simpledialog.askstring("重命名", "请输入新操作名称：", initialvalue=session["name"])
                if new_name and new_name.strip():
                    session["name"] = new_name.strip()
                    self.save_session(session)
                    self.update_history_list()
                    messagebox.showinfo("成功", "操作名称已更新")
                break

    def delete_session(self):
        selected_item = self.history_tree.selection()
        if not selected_item:
            messagebox.showwarning("提示", "请先选择一条历史记录")
            return
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                if messagebox.askyesno("确认删除", f"确定删除「{session['name']}」及所有截图？\n此操作不可恢复！"):
                    # 删除JSON配置文件
                    session_path = os.path.join(self.sessions_dir, f"{session_id}.json")
                    if os.path.exists(session_path):
                        try:
                            os.remove(session_path)
                        except Exception as e:
                            messagebox.showerror("错误", f"删除配置文件失败：{str(e)}")
                            return

                    # 删除截图文件夹
                    images_dir = os.path.join(self.sessions_dir, session_id)
                    if os.path.exists(images_dir):
                        try:
                            for img_file in os.listdir(images_dir):
                                os.remove(os.path.join(images_dir, img_file))
                            os.rmdir(images_dir)
                        except Exception as e:
                            messagebox.showerror("错误", f"删除截图文件失败：{str(e)}")
                            return

                    # 更新列表
                    self.history_sessions.remove(session)
                    self.update_history_list()
                    messagebox.showinfo("成功", "历史记录已删除")
                break

    def start_capture(self):
        if self.is_capturing:
            messagebox.showwarning("提示", "捕捉已在运行中，无需重复启动")
            return

        # 获取会话名称
        session_name = simpledialog.askstring("操作名称", "请输入此次捕捉的操作名称：")
        if not session_name or not session_name.strip():
            messagebox.showwarning("提示", "操作名称不能为空")
            return
        session_name = session_name.strip()

        # 初始化会话数据
        self.current_session_id = str(uuid.uuid4())
        self.current_session = {
            "id": self.current_session_id,
            "name": session_name,
            "description": simpledialog.askstring("操作描述", "请输入此次捕捉的描述（可选）：") or "",
            "start_time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "end_time": "",
            "duration": 0,
            "captures": []
        }

        # 创建截图存储目录
        self.images_dir = os.path.join(self.sessions_dir, self.current_session_id)
        try:
            os.makedirs(self.images_dir, exist_ok=True)
        except Exception as e:
            messagebox.showerror("目录错误", f"创建截图目录失败：{str(e)}")
            return

        # 注册快捷键（修复热键注册逻辑）
        try:
            # 保存热键对象用于后续移除
            self.hotkey_obj = keyboard.add_hotkey(self.hotkey, self.capture_screen)
        except PermissionError:
            messagebox.showerror("权限错误",
                                 "快捷键注册失败！\n请以管理员身份运行程序（Windows）或授予键盘访问权限（macOS）")
            return
        except Exception as e:
            messagebox.showerror("错误", f"快捷键注册失败：{str(e)}")
            return

        # 启动捕捉状态
        self.is_capturing = True
        self.start_time = time.time()
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.root.title(f"屏幕捕捉工具 - 捕捉中（{self.hotkey.upper()}截图）")
        self.status_label.config(text=f"状态：捕捉中（{self.hotkey.upper()}截图）", foreground="#d32f2f")

        # 最小化主窗口并提示
        self.root.iconify()
        messagebox.showinfo("捕捉启动",
                            f"捕捉已开始！\n• 按 {self.hotkey.upper()} 截取当前屏幕\n• 点击「停止捕捉」结束操作\n• 截图自动保存至：{self.images_dir}")

    def stop_capture(self):
        if not self.is_capturing:
            return

        # 清理快捷键（修复热键移除逻辑）
        if self.hotkey_obj:
            try:
                keyboard.remove_hotkey(self.hotkey_obj)
                self.hotkey_obj = None
            except Exception as e:
                messagebox.showwarning("警告", f"快捷键清理失败：{str(e)}")

        # 更新会话数据
        self.is_capturing = False
        self.end_time = time.time()
        self.current_session["end_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.current_session["duration"] = int(self.end_time - self.start_time)

        # 恢复界面状态
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        self.root.title("屏幕捕捉工具 - 就绪")
        self.status_label.config(text="状态：就绪", foreground="#666")

        # 保存会话并刷新历史
        self.save_current_session()
        self.load_history_sessions()
        self.update_history_list()

        # 打开编辑窗口
        if self.current_session["captures"]:
            self.open_editor_window(self.current_session)
        else:
            messagebox.showinfo("提示", "此次捕捉未生成任何截图，无需编辑")

    def capture_screen(self):
        if not self.is_capturing:
            return

        try:
            screenshot = ImageGrab.grab()
        except ImportError:
            messagebox.showerror("依赖缺失",
                                 "Linux 系统需安装 python3-xlib 才能截图\n执行命令：sudo apt install python3-xlib")
            self.stop_capture()
            return
        except Exception as e:
            messagebox.showerror("截图错误", f"截图失败：{str(e)}")
            return

        # 保存截图
        capture_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        img_filename = f"capture_{capture_time}.png"
        img_path = os.path.join(self.images_dir, img_filename)
        try:
            screenshot.save(img_path)
        except Exception as e:
            messagebox.showerror("保存错误", f"截图保存失败：{str(e)}")
            return

        # 添加到会话记录
        capture_count = len(self.current_session["captures"]) + 1
        self.current_session["captures"].append({
            "id": capture_count,
            "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "description": f"第{capture_count}次记录",
            "image_path": img_path
        })

        # 显示临时提示
        self.show_temp_tip(f"已完成第{capture_count}次截图")

    def show_temp_tip(self, message):
        tip_window = tk.Toplevel(self.root)
        tip_window.overrideredirect(True)
        # 优化提示窗口位置计算
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = len(message) * 10 + 40
        window_height = 50
        x_pos = max(0, screen_width - window_width - 20)
        y_pos = max(0, screen_height - window_height - 40)
        tip_window.geometry(f"{window_width}x{window_height}+{x_pos}+{y_pos}")
        tip_window.configure(bg="#2196f3")
        tip_window.attributes("-topmost", True)

        ttk.Label(tip_window, text=message, background="#2196f3", foreground="white", padding=10).pack(fill=tk.BOTH,
                                                                                                       expand=True)
        tip_window.after(2000, tip_window.destroy)

    def open_editor_window(self, session):
        editor_window = tk.Toplevel(self.root)
        editor_window.title(f"编辑操作记录 - {session['name']}")
        editor_window.geometry("900x700")
        editor_window.resizable(True, True)
        editor_window.transient(self.root)

        main_frame = ttk.Frame(editor_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 会话基本信息
        info_frame = ttk.LabelFrame(main_frame, text="操作基本信息", padding="10")
        info_frame.pack(fill=tk.X, pady=5)

        # 名称输入
        ttk.Label(info_frame, text="操作名称：").grid(row=0, column=0, sticky=tk.W, pady=3)
        name_var = tk.StringVar(value=session["name"])
        name_entry = ttk.Entry(info_frame, textvariable=name_var, width=60)
        name_entry.grid(row=0, column=1, sticky=tk.W, pady=3)

        # 描述输入
        ttk.Label(info_frame, text="操作描述：").grid(row=1, column=0, sticky=tk.NW, pady=3)
        desc_text = scrolledtext.ScrolledText(info_frame, width=60, height=4, wrap=tk.WORD)
        desc_text.insert(tk.END, session["description"])
        desc_text.grid(row=1, column=1, sticky=tk.W, pady=3)

        # 只读信息
        info_labels = [
            ("开始时间：", session["start_time"]),
            ("结束时间：", session["end_time"]),
            ("操作时长：", DocxReportGenerator.format_duration(session["duration"])),
            ("截图数量：", str(len(session["captures"])))
        ]
        for i, (label_text, value) in enumerate(info_labels, start=2):
            ttk.Label(info_frame, text=label_text).grid(row=i, column=0, sticky=tk.W, pady=3)
            ttk.Label(info_frame, text=value).grid(row=i, column=1, sticky=tk.W, pady=3)

        # 保存信息按钮
        def save_session_info():
            session["name"] = name_var.get().strip()
            session["description"] = desc_text.get("1.0", tk.END).strip()
            self.save_session(session)
            self.update_history_list()
            messagebox.showinfo("成功", "操作信息已更新")

        ttk.Button(info_frame, text="保存信息", command=save_session_info).grid(row=0, column=2, padx=10, sticky=tk.N)

        # 捕捉记录列表
        ttk.Label(main_frame, text="截图记录（双击预览，点击操作列编辑/删除）", font=("SimHei", 12)).pack(anchor=tk.W,
                                                                                                       pady=5)
        captures_frame = ttk.Frame(main_frame)
        captures_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 截图表格
        columns = ("id", "time", "description", "action")
        captures_tree = ttk.Treeview(captures_frame, columns=columns, show="headings", height=12)
        captures_tree.heading("id", text="序号")
        captures_tree.heading("time", text="截图时间")
        captures_tree.heading("description", text="描述")
        captures_tree.heading("action", text="操作")
        captures_tree.column("id", width=60, anchor=tk.CENTER)
        captures_tree.column("time", width=180)
        captures_tree.column("description", width=400)
        captures_tree.column("action", width=120, anchor=tk.CENTER)
        captures_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 滚动条
        scrollbar = ttk.Scrollbar(captures_frame, orient=tk.VERTICAL, command=captures_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        captures_tree.configure(yscrollcommand=scrollbar.set)

        # 加载截图数据
        for capture in session["captures"]:
            captures_tree.insert("", tk.END, values=(
                capture["id"],
                capture["time"],
                capture["description"],
                "编辑 | 删除"
            ), tags=(capture["id"],))

        # 截图列表操作
        def handle_capture_operation(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region != "cell":
                return
            row = captures_tree.identify_row(event.y)
            col = captures_tree.identify_column(event.x)
            if col != "#4":
                return

            capture_id = int(captures_tree.item(row)["tags"][0])
            capture_idx = None
            target_capture = None
            for i, cap in enumerate(session["captures"]):
                if cap["id"] == capture_id:
                    capture_idx = i
                    target_capture = cap
                    break
            if not target_capture:
                return

            x, y, width, height = captures_tree.bbox(row, col)
            if event.x < x + width / 2:
                # 编辑描述
                new_desc = simpledialog.askstring(
                    f"编辑第{capture_id}次截图描述",
                    "请输入新描述：",
                    initialvalue=target_capture["description"]
                )
                if new_desc is not None:
                    target_capture["description"] = new_desc
                    captures_tree.item(row, values=(
                        capture_id,
                        target_capture["time"],
                        new_desc,
                        "编辑 | 删除"
                    ))
                    self.save_session(session)
            else:
                # 删除截图
                if messagebox.askyesno("确认删除", f"确定删除第{capture_id}次截图？\n此操作不可恢复！"):
                    try:
                        if os.path.exists(target_capture["image_path"]):
                            os.remove(target_capture["image_path"])
                    except Exception as e:
                        messagebox.showerror("错误", f"删除图片失败：{str(e)}")
                        return

                    del session["captures"][capture_idx]
                    for i, cap in enumerate(session["captures"]):
                        cap["id"] = i + 1

                    # 刷新表格
                    for item in captures_tree.get_children():
                        captures_tree.delete(item)
                    for cap in session["captures"]:
                        captures_tree.insert("", tk.END, values=(
                            cap["id"],
                            cap["time"],
                            cap["description"],
                            "编辑 | 删除"
                        ), tags=(cap["id"],))
                    self.save_session(session)

        captures_tree.bind("<Button-1>", handle_capture_operation)

        # 双击预览截图
        def preview_screenshot(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region != "cell" or captures_tree.identify_column(event.x) == "#4":
                return
            row = captures_tree.identify_row(event.y)
            capture_id = int(captures_tree.item(row)["tags"][0])
            target_capture = None
            for cap in session["captures"]:
                if cap["id"] == capture_id:
                    target_capture = cap
                    break
            if not target_capture or not os.path.exists(target_capture["image_path"]):
                messagebox.showwarning("提示", "截图文件已损坏或不存在")
                return

            # 预览窗口
            preview_win = tk.Toplevel(editor_window)
            preview_win.title(f"截图预览 - 第{capture_id}次记录")
            preview_win.geometry("800x600")

            # 滚动条容器
            canvas_frame = ttk.Frame(preview_win)
            canvas_frame.pack(fill=tk.BOTH, expand=True)
            x_scroll = ttk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
            y_scroll = ttk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
            canvas = tk.Canvas(
                canvas_frame,
                xscrollcommand=x_scroll.set,
                yscrollcommand=y_scroll.set,
                bg="#f5f5f5"
            )
            x_scroll.config(command=canvas.xview)
            y_scroll.config(command=canvas.yview)
            x_scroll.pack(side=tk.BOTTOM, fill=tk.X)
            y_scroll.pack(side=tk.RIGHT, fill=tk.Y)
            canvas.pack(fill=tk.BOTH, expand=True)

            # 加载并显示图片（优化缩放逻辑）
            try:
                img = Image.open(target_capture["image_path"])
                img_width, img_height = img.size
                # 计算最佳缩放比例（考虑窗口和屏幕大小）
                screen_width = preview_win.winfo_screenwidth()
                screen_height = preview_win.winfo_screenheight()
                max_width = min(800, screen_width - 100)
                max_height = min(600, screen_height - 100)
                scale = min(max_width / img_width, max_height / img_height, 1.0)
                new_width = int(img_width * scale)
                new_height = int(img_height * scale)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img)
                canvas.create_image(0, 0, anchor=tk.NW, image=photo)
                canvas.image = photo
                canvas.config(scrollregion=(0, 0, new_width, new_height))
            except Exception as e:
                messagebox.showerror("预览错误", f"图片加载失败：{str(e)}")
                preview_win.destroy()
                return

            ttk.Label(preview_win, text=f"描述：{target_capture['description']}", font=("SimHei", 10)).pack(pady=5)

        captures_tree.bind("<Double-1>", preview_screenshot)

        # 报告生成
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)

        def generate_report():
            if not session["captures"]:
                messagebox.showwarning("提示", "无截图记录，无法生成报告")
                return

            # 选择报告格式
            format_win = tk.Toplevel(editor_window)
            format_win.title("选择报告格式")
            format_win.geometry("350x250")
            format_win.transient(editor_window)
            format_win.grab_set()

            ttk.Label(format_win, text="请选择报告输出格式：", font=("SimHei", 12)).pack(pady=20)
            format_var = tk.StringVar(value="docx")

            formats = [
                ("Word 文档 (.docx)", "docx", "兼容性好，支持编辑"),
                ("PDF 文档 (.pdf)", "pdf", "格式固定，跨平台"),
                ("Markdown 文档 (.md)", "md", "轻量文本，支持代码块")
            ]
            for text, val, desc in formats:
                frame = ttk.Frame(format_win)
                frame.pack(anchor=tk.W, padx=50, pady=3)
                ttk.Radiobutton(frame, text=text, variable=format_var, value=val).pack(side=tk.LEFT)
                ttk.Label(frame, text=desc, font=("SimHei", 8), foreground="#666").pack(side=tk.LEFT, padx=10)

            def confirm_generate():
                report_format = format_var.get()
                format_win.destroy()

                filename = f"{session['name']}_操作报告.{report_format}"
                save_path = filedialog.asksaveasfilename(
                    defaultextension=f".{report_format}",
                    filetypes=[(f"{report_format.upper()} 文档", f"*.{report_format}")],
                    initialfile=filename,
                    title=f"保存{report_format.upper()}报告"
                )
                if not save_path:
                    return

                # 显示进度提示
                progress_win = tk.Toplevel(editor_window)
                progress_win.title("生成中")
                progress_win.geometry("300x100")
                progress_win.transient(editor_window)
                progress_win.grab_set()
                ttk.Label(progress_win, text=f"正在生成{report_format.upper()}报告...", font=("SimHei", 12)).pack(
                    expand=True)
                progress_win.update()

                try:
                    if report_format == "docx":
                        success = DocxReportGenerator.generate_docx_report(session, save_path)
                        if success:
                            messagebox.showinfo("成功", f"DOCX报告已生成：\n{save_path}")
                    elif report_format == "pdf":
                        self.generate_pdf_report(session, save_path)
                        messagebox.showinfo("成功", f"PDF报告已生成：\n{save_path}")
                    elif report_format == "md":
                        self.generate_md_report(session, save_path, ask_relative=True)
                        messagebox.showinfo("成功", f"MD报告已生成：\n{save_path}")

                    if messagebox.askyesno("打开报告", "是否立即打开生成的报告？"):
                        os.startfile(save_path)
                except Exception as e:
                    messagebox.showerror("生成失败", f"报告生成错误：{str(e)}")
                finally:
                    progress_win.destroy()

            ttk.Button(format_win, text="确认生成", command=confirm_generate).pack(pady=15)

        ttk.Button(btn_frame, text="生成操作报告", command=generate_report).pack(side=tk.LEFT, padx=5)

        # 作废操作
        def discard_operation():
            if messagebox.askyesno("确认作废", f"确定作废「{session['name']}」操作记录？\n所有截图将被删除，不可恢复！"):
                self.delete_session_by_id(session["id"])
                editor_window.destroy()

        ttk.Button(btn_frame, text="作废此操作", command=discard_operation).pack(side=tk.LEFT, padx=5)

    def generate_pdf_report(self, session, save_path):
        # 优化字体查找逻辑，增加更多备选字体
        font_candidates = []
        if sys.platform == "win32":
            font_candidates = [
                "C:/Windows/Fonts/simsun.ttc",  # 宋体
                "C:/Windows/Fonts/simhei.ttf",  # 黑体
                "C:/Windows/Fonts/microsoftyahei.ttf"  # 微软雅黑
            ]
        elif sys.platform == "darwin":
            font_candidates = [
                "/Library/Fonts/Songti.ttc",  # 宋体
                "/Library/Fonts/Heiti TC.ttc",  # 黑体
                "/Library/Fonts/Microsoft YaHei.ttc"  # 微软雅黑
            ]
        elif sys.platform == "linux":
            font_candidates = [
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # 文泉驿正黑
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",  # Noto Sans
                "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"  # Droid Sans
            ]

        # 查找可用字体
        font_path = None
        for candidate in font_candidates:
            if os.path.exists(candidate):
                font_path = candidate
                break

        if not font_path:
            raise FileNotFoundError(
                f"未找到可用中文字体文件，请安装字体后重试。\n搜索路径：\n" + "\n".join(font_candidates)
            )

        pdfmetrics.registerFont(TTFont('Chinese', font_path))
        pdfmetrics.registerFontFamily('Chinese', normal='Chinese')

        doc = SimpleDocTemplate(save_path, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72,
                                bottomMargin=72)
        styles = getSampleStyleSheet()

        # 自定义中文样式
        styles.add(ParagraphStyle(
            name='ChineseTitle',
            fontName='Chinese',
            fontSize=24,
            spaceAfter=20,
            alignment=1
        ))
        styles.add(ParagraphStyle(
            name='ChineseHeading1',
            fontName='Chinese',
            fontSize=16,
            spaceAfter=12,
            spaceBefore=12,
            bold=True
        ))
        styles.add(ParagraphStyle(
            name='ChineseHeading2',
            fontName='Chinese',
            fontSize=14,
            spaceAfter=8,
            spaceBefore=8,
            bold=True
        ))
        styles.add(ParagraphStyle(
            name='ChineseNormal',
            fontName='Chinese',
            fontSize=12,
            spaceAfter=6,
            leading=18
        ))
        styles.add(ParagraphStyle(
            name='ChineseCaption',
            fontName='Chinese',
            fontSize=10,
            spaceAfter=12,
            italic=True,
            alignment=1
        ))

        # 构建内容
        elements = []
        elements.append(Paragraph("操作记录报告", styles['ChineseTitle']))
        elements.append(
            Paragraph(f"报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        # 基本信息
        elements.append(Paragraph("一、操作基本信息", styles['ChineseHeading1']))
        elements.append(Spacer(1, 8))
        info_items = [
            f"操作名称：{session['name']}",
            f"操作描述：{session['description'] or '无'}",
            f"开始时间：{session['start_time']}",
            f"结束时间：{session['end_time']}",
            f"操作时长：{DocxReportGenerator.format_duration(session['duration'])}"
        ]
        for item in info_items:
            elements.append(Paragraph(item, styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        # 操作步骤
        elements.append(Paragraph("二、详细操作步骤", styles['ChineseHeading1']))
        elements.append(Paragraph(f"共 {len(session['captures'])} 步操作", styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        for i, capture in enumerate(session["captures"], 1):
            elements.append(Paragraph(f"步骤 {i}", styles['ChineseHeading2']))
            elements.append(Paragraph(f"描述：{capture['description']}", styles['ChineseNormal']))
            elements.append(Paragraph(f"截图时间：{capture['time']}", styles['ChineseNormal']))
            elements.append(Spacer(1, 8))

            # 插入截图
            try:
                if os.path.exists(capture["image_path"]):
                    img = Image.open(capture["image_path"])
                    img_width, img_height = img.size
                    scale = min(400 / img_width, 1.0)
                    new_width = int(img_width * scale)
                    new_height = int(img_height * scale)
                    pdf_img = RLImage(capture["image_path"], width=new_width, height=new_height)
                    elements.append(pdf_img)
                    elements.append(Paragraph(f"图 {i}：步骤 {i} 截图", styles['ChineseCaption']))
                else:
                    elements.append(Paragraph(f"[截图文件不存在：{capture['image_path']}]", styles['ChineseNormal']))
            except Exception as e:
                elements.append(Paragraph(f"[截图加载失败：{str(e)}]", styles['ChineseNormal']))

            if i != len(session["captures"]):
                elements.append(Spacer(1, 24))
                elements.append(Paragraph("-" * 60, styles['ChineseNormal']))
                elements.append(Spacer(1, 24))

        doc.build(elements)

    def generate_md_report(self, session, save_path, ask_relative=False):
        """补全Markdown报告生成功能"""
        use_relative = False
        img_relative_dir = ""

        if ask_relative:
            # 询问是否使用相对路径
            if messagebox.askyesno(
                    "路径选择",
                    "是否使用相对路径存储图片引用？\n（适合报告与图片一起移动的场景）"
            ):
                use_relative = True
                # 创建图片存放目录（与报告同目录下的images文件夹）
                report_dir = os.path.dirname(save_path)
                img_relative_dir = os.path.join(report_dir, "images")
                os.makedirs(img_relative_dir, exist_ok=True)

        # 生成Markdown内容
        md_content = []
        md_content.append(f"# 操作记录报告：{session['name']}")
        md_content.append(f"> 报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        md_content.append("")

        # 基本信息
        md_content.append("## 一、基本信息")
        md_content.append(f"- 操作名称：{session['name']}")
        md_content.append(f"- 操作描述：{session['description'] or '无'}")
        md_content.append(f"- 开始时间：{session['start_time']}")
        md_content.append(f"- 结束时间：{session['end_time']}")
        md_content.append(f"- 操作时长：{DocxReportGenerator.format_duration(session['duration'])}")
        md_content.append(f"- 截图数量：{len(session['captures'])}")
        md_content.append("")

        # 操作步骤
        md_content.append("## 二、详细操作步骤")
        for i, capture in enumerate(session["captures"], 1):
            md_content.append(f"### 步骤 {i}")
            md_content.append(f"- 描述：{capture['description']}")
            md_content.append(f"- 截图时间：{capture['time']}")

            # 处理图片路径
            if use_relative:
                # 复制图片到相对目录
                img_filename = os.path.basename(capture["image_path"])
                dest_path = os.path.join(img_relative_dir, img_filename)
                try:
                    with open(capture["image_path"], 'rb') as src, open(dest_path, 'wb') as dst:
                        dst.write(src.read())
                    img_path = f"images/{img_filename}"
                except Exception as e:
                    md_content.append(f"- [图片复制失败：{str(e)}]")
                    continue
            else:
                img_path = capture["image_path"]

            md_content.append(f"![步骤 {i} 截图]({img_path})")
            md_content.append("")

        # 保存Markdown文件
        try:
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_content))
        except Exception as e:
            raise IOError(f"保存Markdown文件失败：{str(e)}")

    def save_session(self, session):
        """保存会话数据到JSON文件"""
        if not session.get("id"):
            return False
        session_path = os.path.join(self.sessions_dir, f"{session['id']}.json")
        try:
            with open(session_path, 'w', encoding='utf-8') as f:
                json.dump(session, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            messagebox.showerror("保存错误", f"会话保存失败：{str(e)}")
            return False

    def save_current_session(self):
        """保存当前会话"""
        if self.current_session_id and self.current_session:
            self.save_session(self.current_session)

    def load_history_sessions(self):
        """加载历史会话"""
        self.history_sessions = []
        try:
            for filename in os.listdir(self.sessions_dir):
                if filename.endswith(".json"):
                    session_id = filename[:-5]
                    session_path = os.path.join(self.sessions_dir, filename)
                    try:
                        with open(session_path, 'r', encoding='utf-8') as f:
                            session = json.load(f)
                            # 验证会话结构
                            if all(key in session for key in ["id", "name", "start_time", "captures"]):
                                self.history_sessions.append(session)
                    except json.JSONDecodeError:
                        messagebox.showwarning("加载警告", f"会话文件损坏：{filename}")
                    except Exception as e:
                        messagebox.showwarning("加载警告", f"加载会话失败：{str(e)}")
            # 按时间排序（最新的在前面）
            self.history_sessions.sort(key=lambda x: x["start_time"], reverse=True)
        except Exception as e:
            messagebox.showerror("加载错误", f"加载历史记录失败：{str(e)}")

    def update_history_list(self):
        """更新历史记录表格"""
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)
        for session in self.history_sessions:
            self.history_tree.insert("", tk.END, values=(
                session["name"],
                session["start_time"],
                DocxReportGenerator.format_duration(session["duration"]),
                len(session["captures"])
            ), tags=(session["id"],))

    def open_session(self, event):
        """打开选中的历史会话"""
        selected_item = self.history_tree.selection()
        if not selected_item:
            return
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                self.open_editor_window(session)
                break

    def delete_session_by_id(self, session_id):
        """通过ID删除会话"""
        for session in self.history_sessions:
            if session["id"] == session_id:
                self.delete_session()
                break

    def on_main_window_close(self):
        """主窗口关闭时清理资源"""
        # 确保移除热键
        if self.hotkey_obj:
            try:
                keyboard.remove_hotkey(self.hotkey_obj)
            except:
                pass
        self.root.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = ScreenCaptureTool(root)
    root.mainloop()