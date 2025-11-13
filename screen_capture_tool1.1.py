import datetime
import json
import os
import sys
import time
import tkinter as tk
import uuid
from tkinter import ttk, filedialog, messagebox, simpledialog, scrolledtext

import keyboard
from PIL import Image, ImageTk, ImageGrab
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image as RLImage, Spacer

# 确保中文显示正常（ matplotlib 仅用于字体配置兼容，非必需可注释）
try:
    import matplotlib

    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]
except ImportError:
    pass  # 不强制依赖 matplotlib


class ScreenCaptureTool:
    def __init__(self, root):
        self.root = root
        self.root.title("屏幕捕捉工具 - 就绪")
        self.root.geometry("600x400")
        self.root.resizable(True, True)

        # 数据存储配置
        self.current_session_id = None
        self.current_session = {
            "name": "",
            "description": "",
            "start_time": "",
            "end_time": "",
            "duration": 0,
            "captures": []
        }
        self.history_sessions = []
        self.sessions_dir = os.path.join(os.path.expanduser("~"), ".screen_capture_sessions")
        os.makedirs(self.sessions_dir, exist_ok=True)

        # 状态变量
        self.is_capturing = False
        self.hotkey = "ctrl+alt+o"
        self.hotkey_registered = False  # 标记快捷键是否注册成功

        # 加载历史记录（带错误提示）
        self.load_history_sessions()

        # 绑定主窗口关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_main_window_close)

        # 创建主界面
        self.create_main_interface()

    def create_main_interface(self):
        """创建主界面布局"""
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题与状态提示
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=5)
        ttk.Label(title_frame, text="屏幕捕捉工具", font=("SimHei", 16)).pack(side=tk.LEFT)
        self.status_label = ttk.Label(title_frame, text="状态：就绪", font=("SimHei", 10), foreground="#666")
        self.status_label.pack(side=tk.RIGHT)

        # 控制按钮（开始/停止）
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(pady=10)
        self.start_btn = ttk.Button(control_frame, text="开始捕捉", command=self.start_capture)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        self.stop_btn = ttk.Button(control_frame, text="停止捕捉", command=self.stop_capture, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=5)

        # 历史记录区域
        ttk.Label(main_frame, text="历史记录", font=("SimHei", 12)).pack(anchor=tk.W, pady=5)
        history_frame = ttk.Frame(main_frame)
        history_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 历史记录表格
        columns = ("name", "date", "duration", "captures")
        self.history_tree = ttk.Treeview(history_frame, columns=columns, show="headings")
        self.history_tree.heading("name", text="操作名称")
        self.history_tree.heading("date", text="开始时间")
        self.history_tree.heading("duration", text="时长")
        self.history_tree.heading("captures", text="截图数")
        self.history_tree.column("name", width=180)
        self.history_tree.column("date", width=180)
        self.history_tree.column("duration", width=80, anchor=tk.CENTER)
        self.history_tree.column("captures", width=80, anchor=tk.CENTER)
        self.history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 滚动条
        scrollbar = ttk.Scrollbar(history_frame, orient=tk.VERTICAL, command=self.history_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.history_tree.configure(yscrollcommand=scrollbar.set)

        # 历史记录右键菜单
        self.history_menu = tk.Menu(self.root, tearoff=0)
        self.history_menu.add_command(label="重命名", command=self.rename_session)
        self.history_menu.add_command(label="删除", command=self.delete_session)

        # 绑定事件
        self.history_tree.bind("<Double-1>", self.open_session)
        self.history_tree.bind("<Button-3>", self.show_history_menu)

        # 更新历史记录列表
        self.update_history_list()

    def show_history_menu(self, event):
        """显示历史记录右键菜单"""
        item = self.history_tree.identify_row(event.y)
        if item:
            self.history_tree.selection_set(item)
            self.history_menu.post(event.x_root, event.y_root)

    def rename_session(self):
        """重命名历史会话"""
        selected_item = self.history_tree.selection()
        if not selected_item:
            messagebox.showwarning("提示", "请先选择一条历史记录")
            return
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                new_name = simpledialog.askstring("重命名", "请输入新操作名称：", initialvalue=session["name"])
                if new_name and new_name.strip():
                    session["name"] = new_name.strip()
                    self.save_session(session)
                    self.update_history_list()
                    messagebox.showinfo("成功", "操作名称已更新")
                break

    def delete_session(self):
        """删除历史会话（含截图文件）"""
        selected_item = self.history_tree.selection()
        if not selected_item:
            messagebox.showwarning("提示", "请先选择一条历史记录")
            return
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                if messagebox.askyesno("确认删除", f"确定删除「{session['name']}」及所有截图？\n此操作不可恢复！"):
                    # 删除JSON配置文件
                    session_path = os.path.join(self.sessions_dir, f"{session_id}.json")
                    if os.path.exists(session_path):
                        os.remove(session_path)
                    # 删除截图文件夹
                    images_dir = os.path.join(self.sessions_dir, session_id)
                    if os.path.exists(images_dir):
                        for img_file in os.listdir(images_dir):
                            os.remove(os.path.join(images_dir, img_file))
                        os.rmdir(images_dir)
                    # 更新列表
                    self.history_sessions.remove(session)
                    self.update_history_list()
                    messagebox.showinfo("成功", "历史记录已删除")
                break

    def start_capture(self):
        """开始屏幕捕捉（含参数校验与环境检查）"""
        if self.is_capturing:
            messagebox.showwarning("提示", "捕捉已在运行中，无需重复启动")
            return

        # 1. 获取会话基本信息
        session_name = simpledialog.askstring("操作名称", "请输入此次捕捉的操作名称：")
        if not session_name or not session_name.strip():
            messagebox.showwarning("提示", "操作名称不能为空")
            return
        session_name = session_name.strip()

        # 2. 初始化会话数据
        self.current_session_id = str(uuid.uuid4())
        self.current_session = {
            "id": self.current_session_id,
            "name": session_name,
            "description": simpledialog.askstring("操作描述", "请输入此次捕捉的描述（可选）：") or "",
            "start_time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "end_time": "",
            "duration": 0,
            "captures": []
        }

        # 3. 创建截图存储目录
        self.images_dir = os.path.join(self.sessions_dir, self.current_session_id)
        os.makedirs(self.images_dir, exist_ok=True)

        # 4. 注册快捷键（处理权限异常）
        try:
            keyboard.add_hotkey(self.hotkey, self.capture_screen)
            self.hotkey_registered = True
        except PermissionError:
            messagebox.showerror("权限错误",
                                 "快捷键注册失败！\n请以管理员身份运行程序（Windows）或授予键盘访问权限（macOS）")
            return
        except Exception as e:
            messagebox.showerror("错误", f"快捷键注册失败：{str(e)}")
            return

        # 5. 启动捕捉状态
        self.is_capturing = True
        self.start_time = time.time()
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.root.title(f"屏幕捕捉工具 - 捕捉中（{self.hotkey.upper()}截图）")
        self.status_label.config(text=f"状态：捕捉中（{self.hotkey.upper()}截图）", foreground="#d32f2f")

        # 6. 最小化主窗口并提示
        self.root.iconify()
        messagebox.showinfo("捕捉启动",
                            f"捕捉已开始！\n• 按 {self.hotkey.upper()} 截取当前屏幕\n• 点击「停止捕捉」结束操作\n• 截图自动保存至：{self.images_dir}")

    def stop_capture(self):
        """停止屏幕捕捉（含资源清理）"""
        if not self.is_capturing:
            return

        # 1. 清理快捷键
        if self.hotkey_registered:
            try:
                keyboard.remove_hotkey(self.hotkey)
                self.hotkey_registered = False
            except KeyError:
                pass  # 快捷键未注册，忽略错误
            except Exception as e:
                messagebox.showwarning("警告", f"快捷键清理失败：{str(e)}")

        # 2. 更新会话数据
        self.is_capturing = False
        self.end_time = time.time()
        self.current_session["end_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.current_session["duration"] = int(self.end_time - self.start_time)

        # 3. 恢复界面状态
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        self.root.title("屏幕捕捉工具 - 就绪")
        self.status_label.config(text="状态：就绪", foreground="#666")

        # 4. 保存会话并刷新历史
        self.save_current_session()
        self.load_history_sessions()
        self.update_history_list()

        # 5. 打开编辑窗口
        if self.current_session["captures"]:
            self.open_editor_window(self.current_session)
        else:
            messagebox.showinfo("提示", "此次捕捉未生成任何截图，无需编辑")

    def capture_screen(self):
        """执行屏幕截图（跨平台兼容处理）"""
        if not self.is_capturing:
            return

        try:
            # 跨平台截图：Windows/macOS 直接使用 ImageGrab，Linux 需要 python3-xlib
            screenshot = ImageGrab.grab()
        except ImportError:
            messagebox.showerror("依赖缺失",
                                 "Linux 系统需安装 python3-xlib 才能截图\n执行命令：sudo apt install python3-xlib")
            self.stop_capture()
            return
        except Exception as e:
            messagebox.showerror("截图错误", f"截图失败：{str(e)}")
            return

        # 保存截图
        capture_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        img_filename = f"capture_{capture_time}.png"
        img_path = os.path.join(self.images_dir, img_filename)
        try:
            screenshot.save(img_path)
        except Exception as e:
            messagebox.showerror("保存错误", f"截图保存失败：{str(e)}")
            return

        # 添加到会话记录
        capture_count = len(self.current_session["captures"]) + 1
        self.current_session["captures"].append({
            "id": capture_count,
            "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "description": f"第{capture_count}次记录",
            "image_path": img_path
        })

        # 显示临时提示
        self.show_temp_tip(f"已完成第{capture_count}次截图")

    def show_temp_tip(self, message):
        """显示临时提示窗口（2秒后自动关闭）"""
        tip_window = tk.Toplevel(self.root)
        tip_window.overrideredirect(True)
        tip_window.geometry(
            f"{len(message) * 10 + 40}x50+{self.root.winfo_screenwidth() - 300}+{self.root.winfo_screenheight() - 120}")
        tip_window.configure(bg="#2196f3")
        tip_window.attributes("-topmost", True)

        # 提示文本（白色）
        ttk.Label(tip_window, text=message, background="#2196f3", foreground="white", padding=10).pack(fill=tk.BOTH,
                                                                                                       expand=True)

        # 2秒后关闭
        tip_window.after(2000, tip_window.destroy)

    def open_editor_window(self, session):
        """打开会话编辑窗口（优化多行描述与截图预览）"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title(f"编辑操作记录 - {session['name']}")
        editor_window.geometry("900x700")
        editor_window.resizable(True, True)
        editor_window.transient(self.root)  # 依附主窗口

        main_frame = ttk.Frame(editor_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 1. 会话基本信息（优化描述输入为多行）
        info_frame = ttk.LabelFrame(main_frame, text="操作基本信息", padding="10")
        info_frame.pack(fill=tk.X, pady=5)

        # 名称输入
        ttk.Label(info_frame, text="操作名称：").grid(row=0, column=0, sticky=tk.W, pady=3)
        name_var = tk.StringVar(value=session["name"])
        name_entry = ttk.Entry(info_frame, textvariable=name_var, width=60)
        name_entry.grid(row=0, column=1, sticky=tk.W, pady=3)

        # 描述输入（多行Text）
        ttk.Label(info_frame, text="操作描述：").grid(row=1, column=0, sticky=tk.NW, pady=3)
        desc_text = scrolledtext.ScrolledText(info_frame, width=60, height=4, wrap=tk.WORD)
        desc_text.insert(tk.END, session["description"])
        desc_text.grid(row=1, column=1, sticky=tk.W, pady=3)

        # 只读信息（时间、时长等）
        info_labels = [
            ("开始时间：", session["start_time"]),
            ("结束时间：", session["end_time"]),
            ("操作时长：", self.format_duration(session["duration"])),
            ("截图数量：", str(len(session["captures"])))
        ]
        for i, (label_text, value) in enumerate(info_labels, start=2):
            ttk.Label(info_frame, text=label_text).grid(row=i, column=0, sticky=tk.W, pady=3)
            ttk.Label(info_frame, text=value).grid(row=i, column=1, sticky=tk.W, pady=3)

        # 保存信息按钮
        def save_session_info():
            session["name"] = name_var.get().strip()
            session["description"] = desc_text.get("1.0", tk.END).strip()
            self.save_session(session)
            self.update_history_list()
            messagebox.showinfo("成功", "操作信息已更新")

        ttk.Button(info_frame, text="保存信息", command=save_session_info).grid(row=0, column=2, padx=10, sticky=tk.N)

        # 2. 捕捉记录列表
        ttk.Label(main_frame, text="截图记录（双击预览，点击操作列编辑/删除）", font=("SimHei", 12)).pack(anchor=tk.W,
                                                                                                       pady=5)
        captures_frame = ttk.Frame(main_frame)
        captures_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 截图表格
        columns = ("id", "time", "description", "action")
        captures_tree = ttk.Treeview(captures_frame, columns=columns, show="headings", height=12)
        captures_tree.heading("id", text="序号")
        captures_tree.heading("time", text="截图时间")
        captures_tree.heading("description", text="描述")
        captures_tree.heading("action", text="操作")
        captures_tree.column("id", width=60, anchor=tk.CENTER)
        captures_tree.column("time", width=180)
        captures_tree.column("description", width=400)
        captures_tree.column("action", width=120, anchor=tk.CENTER)
        captures_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 滚动条
        scrollbar = ttk.Scrollbar(captures_frame, orient=tk.VERTICAL, command=captures_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        captures_tree.configure(yscrollcommand=scrollbar.set)

        # 加载截图数据
        for capture in session["captures"]:
            captures_tree.insert("", tk.END, values=(
                capture["id"],
                capture["time"],
                capture["description"],
                "编辑 | 删除"
            ), tags=(capture["id"],))

        # 3. 截图列表操作（编辑描述/删除）
        def handle_capture_operation(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region != "cell":
                return
            row = captures_tree.identify_row(event.y)
            col = captures_tree.identify_column(event.x)
            if col != "#4":  # 仅操作列响应
                return

            # 获取当前截图记录
            capture_id = int(captures_tree.item(row)["tags"][0])
            capture_idx = None
            target_capture = None
            for i, cap in enumerate(session["captures"]):
                if cap["id"] == capture_id:
                    capture_idx = i
                    target_capture = cap
                    break
            if not target_capture:
                return

            # 计算点击位置（左半区=编辑，右半区=删除）
            x, y, width, height = captures_tree.bbox(row, col)
            if event.x < x + width / 2:
                # 编辑描述
                new_desc = simpledialog.askstring(
                    f"编辑第{capture_id}次截图描述",
                    "请输入新描述：",
                    initialvalue=target_capture["description"]
                )
                if new_desc is not None:  # 取消则不处理
                    target_capture["description"] = new_desc
                    captures_tree.item(row, values=(
                        capture_id,
                        target_capture["time"],
                        new_desc,
                        "编辑 | 删除"
                    ))
                    self.save_session(session)
            else:
                # 删除截图
                if messagebox.askyesno("确认删除", f"确定删除第{capture_id}次截图？\n此操作不可恢复！"):
                    # 删除图片文件
                    if os.path.exists(target_capture["image_path"]):
                        os.remove(target_capture["image_path"])
                    # 移除记录并重新编号
                    del session["captures"][capture_idx]
                    for i, cap in enumerate(session["captures"]):
                        cap["id"] = i + 1
                    # 刷新表格
                    for item in captures_tree.get_children():
                        captures_tree.delete(item)
                    for cap in session["captures"]:
                        captures_tree.insert("", tk.END, values=(
                            cap["id"],
                            cap["time"],
                            cap["description"],
                            "编辑 | 删除"
                        ), tags=(cap["id"],))
                    self.save_session(session)

        captures_tree.bind("<Button-1>", handle_capture_operation)

        # 4. 双击预览截图（支持缩放与滚动）
        def preview_screenshot(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region != "cell" or captures_tree.identify_column(event.x) == "#4":
                return
            row = captures_tree.identify_row(event.y)
            capture_id = int(captures_tree.item(row)["tags"][0])
            target_capture = None
            for cap in session["captures"]:
                if cap["id"] == capture_id:
                    target_capture = cap
                    break
            if not target_capture or not os.path.exists(target_capture["image_path"]):
                messagebox.showwarning("提示", "截图文件已损坏或不存在")
                return

            # 预览窗口
            preview_win = tk.Toplevel(editor_window)
            preview_win.title(f"截图预览 - 第{capture_id}次记录")
            preview_win.geometry("800x600")

            # 滚动条容器
            canvas_frame = ttk.Frame(preview_win)
            canvas_frame.pack(fill=tk.BOTH, expand=True)
            x_scroll = ttk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
            y_scroll = ttk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
            canvas = tk.Canvas(
                canvas_frame,
                xscrollcommand=x_scroll.set,
                yscrollcommand=y_scroll.set,
                bg="#f5f5f5"
            )
            x_scroll.config(command=canvas.xview)
            y_scroll.config(command=canvas.yview)
            x_scroll.pack(side=tk.BOTTOM, fill=tk.X)
            y_scroll.pack(side=tk.RIGHT, fill=tk.Y)
            canvas.pack(fill=tk.BOTH, expand=True)

            # 加载并显示图片
            try:
                img = Image.open(target_capture["image_path"])
                img_width, img_height = img.size
                # 按窗口大小缩放（最大800x600）
                scale = min(800 / img_width, 600 / img_height, 1.0)
                new_width = int(img_width * scale)
                new_height = int(img_height * scale)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(img)
                canvas.create_image(0, 0, anchor=tk.NW, image=photo)
                canvas.image = photo  # 保持引用
                canvas.config(scrollregion=(0, 0, new_width, new_height))
            except Exception as e:
                messagebox.showerror("预览错误", f"图片加载失败：{str(e)}")
                preview_win.destroy()
                return

            # 截图描述标签
            ttk.Label(preview_win, text=f"描述：{target_capture['description']}", font=("SimHei", 10)).pack(pady=5)

        captures_tree.bind("<Double-1>", preview_screenshot)

        # 5. 报告生成与作废按钮
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)

        # 生成报告
        def generate_report():
            if not session["captures"]:
                messagebox.showwarning("提示", "无截图记录，无法生成报告")
                return

            # 选择报告格式
            format_win = tk.Toplevel(editor_window)
            format_win.title("选择报告格式")
            format_win.geometry("350x250")
            format_win.transient(editor_window)
            format_win.grab_set()

            ttk.Label(format_win, text="请选择报告输出格式：", font=("SimHei", 12)).pack(pady=20)
            format_var = tk.StringVar(value="docx")

            # 格式选项（带说明）
            formats = [
                ("Word 文档 (.docx)", "docx", "兼容性好，支持编辑"),
                ("PDF 文档 (.pdf)", "pdf", "格式固定，跨平台"),
                ("Markdown 文档 (.md)", "md", "轻量文本，支持代码块")
            ]
            for text, val, desc in formats:
                frame = ttk.Frame(format_win)
                frame.pack(anchor=tk.W, padx=50, pady=3)
                ttk.Radiobutton(frame, text=text, variable=format_var, value=val).pack(side=tk.LEFT)
                ttk.Label(frame, text=desc, font=("SimHei", 8), foreground="#666").pack(side=tk.LEFT, padx=10)

            # 确认生成
            def confirm_generate():
                report_format = format_var.get()
                format_win.destroy()

                # 选择保存路径
                filename = f"{session['name']}_操作报告.{report_format}"
                save_path = filedialog.asksaveasfilename(
                    defaultextension=f".{report_format}",
                    filetypes=[(f"{report_format.upper()} 文档", f"*.{report_format}")],
                    initialfile=filename,
                    title=f"保存{report_format.upper()}报告"
                )
                if not save_path:
                    return

                # 生成对应格式报告
                try:
                    if report_format == "docx":
                        self.generate_docx_report(session, save_path)
                    elif report_format == "pdf":
                        self.generate_pdf_report(session, save_path)
                    elif report_format == "md":
                        self.generate_md_report(session, save_path, ask_relative=True)
                    messagebox.showinfo("成功", f"报告已生成：\n{save_path}")
                    if messagebox.askyesno("打开报告", "是否立即打开生成的报告？"):
                        os.startfile(save_path)
                except Exception as e:
                    messagebox.showerror("生成失败", f"报告生成错误：{str(e)}")

            ttk.Button(format_win, text="确认生成", command=confirm_generate).pack(pady=15)

        ttk.Button(btn_frame, text="生成操作报告", command=generate_report).pack(side=tk.LEFT, padx=5)

        # 作废操作
        def discard_operation():
            if messagebox.askyesno("确认作废", f"确定作废「{session['name']}」操作记录？\n所有截图将被删除，不可恢复！"):
                self.delete_session_by_id(session["id"])
                editor_window.destroy()

        ttk.Button(btn_frame, text="作废此操作", command=discard_operation, style="danger.TButton").pack(side=tk.LEFT,
                                                                                                         padx=5)

        # 自定义危险按钮样式
        editor_window.style = ttk.Style()
        editor_window.style.configure("danger.TButton", foreground="red")

    def generate_docx_report(self, session, save_path):
        """生成 Word 格式报告（保持原逻辑，优化表格样式）"""
        doc = Document()

        # 标题
        doc.add_heading("操作记录报告", 0)
        doc.add_paragraph(f"报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style="Normal")
        doc.add_paragraph()

        # 基本信息表格（带样式）
        doc.add_heading("一、操作基本信息", level=1)
        info_table = doc.add_table(rows=5, cols=2, style="Table Grid")
        info_table.autofit = False
        info_table.columns[0].width = Inches(1.5)
        info_table.columns[1].width = Inches(4.5)
        # 表头内容
        headers = [
            ("操作名称", session["name"]),
            ("操作描述", session["description"]),
            ("开始时间", session["start_time"]),
            ("结束时间", session["end_time"]),
            ("操作时长", self.format_duration(session["duration"]))
        ]
        for i, (label, value) in enumerate(headers):
            cell1 = info_table.cell(i, 0)
            cell1.text = label
            cell1.paragraphs[0].bold = True
            cell2 = info_table.cell(i, 1)
            cell2.text = value

        # 操作步骤（带截图）
        doc.add_heading("二、详细操作步骤", level=1)
        doc.add_paragraph(f"共 {len(session['captures'])} 步操作，以下为详细记录：", style="Normal")
        doc.add_paragraph()

        for i, capture in enumerate(session["captures"], 1):
            # 步骤标题
            doc.add_heading(f"步骤 {i}", level=2)
            # 步骤描述
            doc.add_paragraph(f"描述：{capture['description']}", style="Normal")
            doc.add_paragraph(f"截图时间：{capture['time']}", style="Normal")
            # 步骤截图
            try:
                doc.add_picture(capture["image_path"], width=Inches(6))
                doc.add_caption(f"图 {i}：步骤 {i} 截图", style="Caption")
            except Exception as e:
                doc.add_paragraph(f"[截图加载失败：{str(e)}]", style="Normal")
            # 分页
            if i != len(session["captures"]):
                doc.add_page_break()

        # 保存文档
        doc.save(save_path)

    def generate_pdf_report(self, session, save_path):
        """生成 PDF 报告（修复中文乱码，优化格式）"""
        # 1. 配置中文字体（适配 Windows/macOS/Linux）
        font_path = ""
        if sys.platform == "win32":
            font_path = "C:/Windows/Fonts/simsun.ttc"  # Windows 宋体
        elif sys.platform == "darwin":
            font_path = "/Library/Fonts/Songti.ttc"  # macOS 宋体
        elif sys.platform == "linux":
            font_path = "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"  # Linux 文泉驿黑体

        # 检查字体文件是否存在
        if not os.path.exists(font_path):
            raise FileNotFoundError(f"未找到中文字体文件：{font_path}\n请安装对应字体或手动指定字体路径")

        # 注册中文字体
        pdfmetrics.registerFont(TTFont('SimSun', font_path))
        pdfmetrics.registerFontFamily('SimSun', normal='SimSun')

        # 2. 创建 PDF 文档
        doc = SimpleDocTemplate(save_path, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72,
                                bottomMargin=72)
        styles = getSampleStyleSheet()

        # 自定义中文样式（覆盖默认样式）
        styles.add(ParagraphStyle(
            name='ChineseTitle',
            fontName='SimSun',
            fontSize=24,
            spaceAfter=20,
            alignment=1  # 居中
        ))
        styles.add(ParagraphStyle(
            name='ChineseHeading1',
            fontName='SimSun',
            fontSize=16,
            spaceAfter=12,
            spaceBefore=12,
            bold=True
        ))
        styles.add(ParagraphStyle(
            name='ChineseHeading2',
            fontName='SimSun',
            fontSize=14,
            spaceAfter=8,
            spaceBefore=8,
            bold=True
        ))
        styles.add(ParagraphStyle(
            name='ChineseNormal',
            fontName='SimSun',
            fontSize=12,
            spaceAfter=6,
            leading=18  # 行高
        ))
        styles.add(ParagraphStyle(
            name='ChineseCaption',
            fontName='SimSun',
            fontSize=10,
            spaceAfter=12,
            italic=True,
            alignment=1
        ))

        # 3. 构建 PDF 内容
        elements = []

        # 标题
        elements.append(Paragraph("操作记录报告", styles['ChineseTitle']))
        elements.append(
            Paragraph(f"报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        # 基本信息
        elements.append(Paragraph("一、操作基本信息", styles['ChineseHeading1']))
        elements.append(Spacer(1, 8))
        info_items = [
            f"操作名称：{session['name']}",
            f"操作描述：{session['description'] or '无'}",
            f"开始时间：{session['start_time']}",
            f"结束时间：{session['end_time']}",
            f"操作时长：{self.format_duration(session['duration'])}"
        ]
        for item in info_items:
            elements.append(Paragraph(item, styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        # 操作步骤
        elements.append(Paragraph("二、详细操作步骤", styles['ChineseHeading1']))
        elements.append(Paragraph(f"共 {len(session['captures'])} 步操作，以下为详细记录：", styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        for i, capture in enumerate(session["captures"], 1):
            elements.append(Paragraph(f"步骤 {i}", styles['ChineseHeading2']))
            elements.append(Paragraph(f"描述：{capture['description']}", styles['ChineseNormal']))
            elements.append(Paragraph(f"截图时间：{capture['time']}", styles['ChineseNormal']))
            elements.append(Spacer(1, 8))

            # 插入截图
            try:
                # 缩放图片（最大宽度 400）
                img = Image.open(capture["image_path"])
                img_width, img_height = img.size
                scale = min(400 / img_width, 1.0)
                new_width = int(img_width * scale)
                new_height = int(img_height * scale)

                # 添加图片到 PDF
                pdf_img = RLImage(capture["image_path"], width=new_width, height=new_height)
                elements.append(pdf_img)
                elements.append(Paragraph(f"图 {i}：步骤 {i} 截图", styles['ChineseCaption']))
            except Exception as e:
                elements.append(Paragraph(f"[截图加载失败：{str(e)}]", styles['ChineseNormal']))

            # 分页（最后一步不分页）
            if i != len(session["captures"]):
                elements.append(Spacer(1, 24))
                elements.append(Paragraph("-" * 60, styles['ChineseNormal']))
                elements.append(Spacer(1, 24))

        # 生成 PDF
        doc.build(elements)

    def generate_md_report(self, session, save_path, ask_relative=False):
        """生成 Markdown 报告（支持相对路径选项）"""
        # 询问是否使用相对路径
        use_relative = False
        if ask_relative:
            if messagebox.askyesno("路径选择",
                                   "是否使用相对路径存储图片？\n• 相对路径：报告与图片文件夹同目录时可正常显示\n• 绝对路径：报告移动后仍可显示（默认）"):
                use_relative = True
                # 创建图片子文件夹（若选择相对路径）
                img_relative_dir = os.path.join(os.path.dirname(save_path), f"{session['id']}_images")
                os.makedirs(img_relative_dir, exist_ok=True)
                # 复制图片到相对路径目录
                for capture in session["captures"]:
                    src_path = capture["image_path"]
                    dst_path = os.path.join(img_relative_dir, os.path.basename(src_path))
                    if not os.path.exists(dst_path):
                        Image.open(src_path).save(dst_path)

        # 写入 MD 内容
        with open(save_path, "w", encoding="utf-8") as f:
            # 标题与基本信息
            f.write(f"# 操作记录报告\n\n")
            f.write(f"**报告生成时间**：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"## 一、操作基本信息\n\n")
            f.write(f"| 项目       | 内容                                   |\n")
            f.write(f"|------------|----------------------------------------|\n")
            f.write(f"| 操作名称   | {session['name']}                      |\n")
            f.write(f"| 操作描述   | {session['description'] or '无'}       |\n")
            f.write(f"| 开始时间   | {session['start_time']}                |\n")
            f.write(f"| 结束时间   | {session['end_time']}                  |\n")
            f.write(f"| 操作时长   | {self.format_duration(session['duration'])} |\n")
            f.write(f"| 截图数量   | {len(session['captures'])} 步          |\n\n")

            # 操作步骤
            f.write(f"## 二、详细操作步骤\n\n")
            f.write(f"共 {len(session['captures'])} 步操作，以下为详细记录：\n\n")

            for i, capture in enumerate(session["captures"], 1):
                f.write(f"### 步骤 {i}\n\n")
                f.write(f"- **描述**：{capture['description']}\n")
                f.write(f"- **截图时间**：{capture['time']}\n\n")

                # 图片路径（绝对/相对）
                if use_relative:
                    img_path = os.path.join(f"./{os.path.basename(img_relative_dir)}",
                                            os.path.basename(capture["image_path"]))
                else:
                    img_path = capture["image_path"]

                # 插入图片（支持 Markdown 图片语法）
                f.write(f"![图 {i}：步骤 {i} 截图]({img_path})\n\n")
                f.write(f"---\n\n")

        # 提示相对路径使用方式
        if use_relative:
            messagebox.showinfo("相对路径提示",
                                f"图片已保存至：\n{img_relative_dir}\n\n使用说明：\n1. 请勿修改图片文件夹名称\n2. 报告与图片文件夹需保持同目录\n3. 移动报告时请同时移动图片文件夹")

    def save_current_session(self):
        """保存当前会话到 JSON 文件"""
        if not self.current_session_id:
            return
        session_path = os.path.join(self.sessions_dir, f"{self.current_session_id}.json")
        with open(session_path, "w", encoding="utf-8") as f:
            json.dump(self.current_session, f, ensure_ascii=False, indent=2)

    def save_session(self, session):
        """保存指定会话到 JSON 文件"""
        session_path = os.path.join(self.sessions_dir, f"{session['id']}.json")
        with open(session_path, "w", encoding="utf-8") as f:
            json.dump(session, f, ensure_ascii=False, indent=2)

    def load_history_sessions(self):
        """加载历史会话（带损坏文件提示）"""
        self.history_sessions = []
        if not os.path.exists(self.sessions_dir):
            return

        for filename in os.listdir(self.sessions_dir):
            if not filename.endswith(".json"):
                continue
            session_id = filename[:-5]
            session_path = os.path.join(self.sessions_dir, filename)
            try:
                with open(session_path, "r", encoding="utf-8") as f:
                    session = json.load(f)
                    # 校验会话数据完整性
                    required_keys = ["id", "name", "start_time", "end_time", "duration", "captures"]
                    if all(key in session for key in required_keys):
                        self.history_sessions.append(session)
                    else:
                        print(f"警告：会话文件 {filename} 缺少必要字段，已跳过")
            except json.JSONDecodeError:
                messagebox.showwarning("文件损坏",
                                       f"历史记录文件 {filename} 已损坏，无法加载\n建议手动删除该文件：{session_path}")
            except Exception as e:
                messagebox.showwarning("加载错误", f"加载会话 {filename} 时出错：{str(e)}")

        # 按开始时间倒序排序
        self.history_sessions.sort(key=lambda x: x["start_time"], reverse=True)

    def update_history_list(self):
        """更新历史记录表格"""
        # 清空现有数据
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)
        # 添加新数据
        for session in self.history_sessions:
            self.history_tree.insert("", tk.END, values=(
                session["name"],
                session["start_time"],
                self.format_duration(session["duration"]),
                len(session["captures"])
            ), tags=(session["id"],))

    def format_duration(self, seconds):
        """格式化时长（秒 → 时:分:秒）"""
        hours, remainder = divmod(seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        if hours > 0:
            return f"{hours}h{minutes}m{seconds}s"
        elif minutes > 0:
            return f"{minutes}m{seconds}s"
        else:
            return f"{seconds}s"

    def open_session(self, event):
        """双击打开历史会话"""
        selected_item = self.history_tree.selection()
        if not selected_item:
            return
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                self.open_editor_window(session)
                break

    def delete_session_by_id(self, session_id):
        """通过 ID 删除会话（供作废功能调用）"""
        for session in self.history_sessions:
            if session["id"] == session_id:
                # 删除文件
                session_path = os.path.join(self.sessions_dir, f"{session_id}.json")
                if os.path.exists(session_path):
                    os.remove(session_path)
                # 删除截图
                images_dir = os.path.join(self.sessions_dir, session_id)
                if os.path.exists(images_dir):
                    for img_file in os.listdir(images_dir):
                        os.remove(os.path.join(images_dir, img_file))
                    os.rmdir(images_dir)
                # 更新列表
                self.history_sessions.remove(session)
                self.update_history_list()
                messagebox.showinfo("成功", "操作记录已作废并删除")
                break

    def on_main_window_close(self):
        """主窗口关闭时清理资源"""
        if self.is_capturing:
            self.stop_capture()
        self.root.destroy()


if __name__ == "__main__":
    # 检查必要依赖
    required_libs = [
        ("PIL", "pillow"),
        ("keyboard", "keyboard"),
        ("docx", "python-docx"),
        ("reportlab", "reportlab")
    ]
    missing_libs = []
    for lib_name, pip_name in required_libs:
        try:
            __import__(lib_name)
        except ImportError:
            missing_libs.append(pip_name)

    if missing_libs:
        messagebox.showerror("依赖缺失", f"缺少必要依赖库，请先安装：\npip install {' '.join(missing_libs)}")
        sys.exit(1)

    # 启动应用
    root = tk.Tk()
    app = ScreenCaptureTool(root)
    root.mainloop()