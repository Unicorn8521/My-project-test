import datetime
import json
import os
import sys
import threading
import time
import tkinter as tk
import uuid
from tkinter import ttk, filedialog, messagebox, simpledialog

import keyboard
# 确保中文显示正常
import matplotlib
from PIL import Image, ImageTk, ImageGrab
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image as RLImage, Spacer

matplotlib.use('Agg')
import matplotlib.pyplot as plt
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]

class ScreenCaptureTool:
    def __init__(self, root):
        self.root = root
        self.root.title("屏幕捕捉工具")
        self.root.geometry("600x400")
        self.root.resizable(True, True)
        
        # 数据存储
        self.current_session_id = None
        self.current_session = {
            "name": "",
            "description": "",
            "start_time": "",
            "end_time": "",
            "duration": 0,
            "captures": []
        }
        self.history_sessions = []
        self.sessions_dir = os.path.join(os.path.expanduser("~"), ".screen_capture_sessions")
        os.makedirs(self.sessions_dir, exist_ok=True)
        
        # 状态变量
        self.is_capturing = False
        self.hotkey = "ctrl+alt+o"
        
        # 加载历史记录
        self.load_history_sessions()
        
        # 创建主界面
        self.create_main_interface()
    
    def create_main_interface(self):
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        ttk.Label(main_frame, text="屏幕捕捉工具", font=("SimHei", 16)).pack(pady=10)
        
        # 控制按钮
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(pady=10)
        
        self.start_btn = ttk.Button(control_frame, text="开始捕捉", command=self.start_capture)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        
        self.stop_btn = ttk.Button(control_frame, text="停止捕捉", command=self.stop_capture, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=5)
        
        # 历史记录区域
        ttk.Label(main_frame, text="历史记录", font=("SimHei", 12)).pack(anchor=tk.W, pady=5)
        
        # 历史记录列表
        history_frame = ttk.Frame(main_frame)
        history_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        columns = ("name", "date", "duration", "captures")
        self.history_tree = ttk.Treeview(history_frame, columns=columns, show="headings")
        
        self.history_tree.heading("name", text="名称")
        self.history_tree.heading("date", text="日期")
        self.history_tree.heading("duration", text="时长")
        self.history_tree.heading("captures", text="截图数量")
        
        self.history_tree.column("name", width=150)
        self.history_tree.column("date", width=150)
        self.history_tree.column("duration", width=80)
        self.history_tree.column("captures", width=80)
        
        self.history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 滚动条
        scrollbar = ttk.Scrollbar(history_frame, orient=tk.VERTICAL, command=self.history_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.history_tree.configure(yscrollcommand=scrollbar.set)
        
        # 右键菜单
        self.history_menu = tk.Menu(self.root, tearoff=0)
        self.history_menu.add_command(label="重命名", command=self.rename_session)
        self.history_menu.add_command(label="删除", command=self.delete_session)
        
        # 绑定事件
        self.history_tree.bind("<Double-1>", self.open_session)
        self.history_tree.bind("<Button-3>", self.show_history_menu)
        
        # 更新历史记录列表
        self.update_history_list()
    
    def show_history_menu(self, event):
        item = self.history_tree.identify_row(event.y)
        if item:
            self.history_tree.selection_set(item)
            self.history_menu.post(event.x_root, event.y_root)
    
    def rename_session(self):
        selected_item = self.history_tree.selection()
        if not selected_item:
            return
            
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                new_name = simpledialog.askstring("重命名", "请输入新名称:", initialvalue=session["name"])
                if new_name and new_name.strip():
                    session["name"] = new_name.strip()
                    self.save_session(session)
                    self.update_history_list()
                break
    
    def delete_session(self):
        selected_item = self.history_tree.selection()
        if not selected_item:
            return
            
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                if messagebox.askyesno("确认删除", f"确定要删除 '{session['name']}' 吗?"):
                    # 删除文件
                    session_path = os.path.join(self.sessions_dir, f"{session_id}.json")
                    if os.path.exists(session_path):
                        os.remove(session_path)
                    
                    # 删除截图
                    images_dir = os.path.join(self.sessions_dir, session_id)
                    if os.path.exists(images_dir):
                        for img_file in os.listdir(images_dir):
                            os.remove(os.path.join(images_dir, img_file))
                        os.rmdir(images_dir)
                    
                    # 更新列表
                    self.history_sessions.remove(session)
                    self.update_history_list()
                break
    
    def start_capture(self):
        # 获取会话名称
        session_name = simpledialog.askstring("会话名称", "请输入此次操作名称:")
        if not session_name or not session_name.strip():
            return
            
        session_desc = simpledialog.askstring("会话描述", "请输入此次操作描述:")
        
        # 初始化会话
        self.current_session_id = str(uuid.uuid4())
        self.current_session = {
            "id": self.current_session_id,
            "name": session_name.strip(),
            "description": session_desc.strip() if session_desc else "",
            "start_time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "end_time": "",
            "duration": 0,
            "captures": []
        }
        
        # 创建截图目录
        self.images_dir = os.path.join(self.sessions_dir, self.current_session_id)
        os.makedirs(self.images_dir, exist_ok=True)
        
        # 开始捕捉
        self.is_capturing = True
        self.start_time = time.time()
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        
        # 注册快捷键
        keyboard.add_hotkey(self.hotkey, self.capture_screen)
        
        # 最小化窗口
        self.root.iconify()
        
        # 显示提示
        messagebox.showinfo("开始捕捉", f"捕捉已开始!\n按 {self.hotkey.upper()} 进行截屏\n点击停止捕捉结束")
    
    def stop_capture(self):
        if not self.is_capturing:
            return
            
        # 停止捕捉
        self.is_capturing = False
        self.end_time = time.time()
        self.current_session["end_time"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.current_session["duration"] = int(self.end_time - self.start_time)
        
        # 移除快捷键
        keyboard.remove_hotkey(self.hotkey)
        
        # 恢复按钮状态
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        
        # 保存会话
        self.save_current_session()
        
        # 刷新历史记录
        self.load_history_sessions()
        self.update_history_list()
        
        # 打开编辑窗口
        self.open_editor_window(self.current_session)
    
    def capture_screen(self):
        if not self.is_capturing:
            return
            
        # 截屏
        screenshot = ImageGrab.grab()
        
        # 保存截图
        capture_time = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        img_filename = f"capture_{capture_time}.png"
        img_path = os.path.join(self.images_dir, img_filename)
        screenshot.save(img_path)
        
        # 添加到会话
        capture_count = len(self.current_session["captures"]) + 1
        self.current_session["captures"].append({
            "id": capture_count,
            "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "description": f"第{capture_count}次记录",
            "image_path": img_path
        })
        
        # 显示提示
        self.show_system_tray_message(f"已完成第{capture_count}次截屏")
    
    def show_system_tray_message(self, message):
        # 创建临时窗口显示提示
        tip_window = tk.Toplevel(self.root)
        tip_window.overrideredirect(True)
        tip_window.geometry("200x50+{}+{}".format(
            self.root.winfo_screenwidth() - 220,
            self.root.winfo_screenheight() - 100
        ))
        tip_window.configure(bg="#f0f0f0")
        
        ttk.Label(tip_window, text=message, background="#f0f0f0", padding=10).pack(fill=tk.BOTH, expand=True)
        
        tip_window.attributes("-topmost", True)
        tip_window.after(2000, tip_window.destroy)
    
    def save_current_session(self):
        if not self.current_session_id:
            return
            
        session_path = os.path.join(self.sessions_dir, f"{self.current_session_id}.json")
        with open(session_path, "w", encoding="utf-8") as f:
            json.dump(self.current_session, f, ensure_ascii=False, indent=2)
    
    def save_session(self, session):
        session_path = os.path.join(self.sessions_dir, f"{session['id']}.json")
        with open(session_path, "w", encoding="utf-8") as f:
            json.dump(session, f, ensure_ascii=False, indent=2)
    
    def load_history_sessions(self):
        self.history_sessions = []
        for filename in os.listdir(self.sessions_dir):
            if filename.endswith(".json"):
                session_id = filename[:-5]
                session_path = os.path.join(self.sessions_dir, filename)
                try:
                    with open(session_path, "r", encoding="utf-8") as f:
                        session = json.load(f)
                        self.history_sessions.append(session)
                except:
                    continue
        
        # 按时间排序
        self.history_sessions.sort(key=lambda x: x["start_time"], reverse=True)
    
    def update_history_list(self):
        # 清空现有项
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)
        
        # 添加历史记录
        for session in self.history_sessions:
            duration = self.format_duration(session["duration"])
            self.history_tree.insert("", tk.END, values=(
                session["name"],
                session["start_time"],
                duration,
                len(session["captures"])
            ), tags=(session["id"],))
    
    def format_duration(self, seconds):
        minutes, seconds = divmod(seconds, 60)
        hours, minutes = divmod(minutes, 60)
        if hours > 0:
            return f"{hours}h{minutes}m"
        elif minutes > 0:
            return f"{minutes}m{seconds}s"
        else:
            return f"{seconds}s"
    
    def open_session(self, event):
        selected_item = self.history_tree.selection()
        if not selected_item:
            return
            
        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                self.open_editor_window(session)
                break
    
    def open_editor_window(self, session):
        # 创建编辑窗口
        editor_window = tk.Toplevel(self.root)
        editor_window.title(f"编辑 - {session['name']}")
        editor_window.geometry("800x600")
        editor_window.resizable(True, True)
        
        # 主框架
        main_frame = ttk.Frame(editor_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 会话信息
        info_frame = ttk.LabelFrame(main_frame, text="会话信息", padding="10")
        info_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(info_frame, text="名称:").grid(row=0, column=0, sticky=tk.W, pady=5)
        name_var = tk.StringVar(value=session["name"])
        ttk.Entry(info_frame, textvariable=name_var, width=50).grid(row=0, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(info_frame, text="描述:").grid(row=1, column=0, sticky=tk.NW, pady=5)
        desc_var = tk.StringVar(value=session["description"])
        ttk.Entry(info_frame, textvariable=desc_var, width=50).grid(row=1, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(info_frame, text="开始时间:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Label(info_frame, text=session["start_time"]).grid(row=2, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(info_frame, text="结束时间:").grid(row=3, column=0, sticky=tk.W, pady=5)
        ttk.Label(info_frame, text=session["end_time"]).grid(row=3, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(info_frame, text="时长:").grid(row=4, column=0, sticky=tk.W, pady=5)
        ttk.Label(info_frame, text=self.format_duration(session["duration"])).grid(row=4, column=1, sticky=tk.W, pady=5)
        
        ttk.Label(info_frame, text="截图数量:").grid(row=5, column=0, sticky=tk.W, pady=5)
        ttk.Label(info_frame, text=str(len(session["captures"]))).grid(row=5, column=1, sticky=tk.W, pady=5)
        
        # 保存信息按钮
        def save_session_info():
            session["name"] = name_var.get().strip()
            session["description"] = desc_var.get().strip()
            self.save_session(session)
            self.load_history_sessions()
            self.update_history_list()
            messagebox.showinfo("保存成功", "会话信息已更新")
        
        ttk.Button(info_frame, text="保存信息", command=save_session_info).grid(row=0, column=2, padx=10)
        
        # 捕捉列表
        ttk.Label(main_frame, text="捕捉记录", font=("SimHei", 12)).pack(anchor=tk.W, pady=5)
        
        # 创建一个框架放置列表和滚动条
        captures_frame = ttk.Frame(main_frame)
        captures_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 捕捉记录列表
        columns = ("id", "time", "description", "action")
        captures_tree = ttk.Treeview(captures_frame, columns=columns, show="headings", height=10)
        
        captures_tree.heading("id", text="序号")
        captures_tree.heading("time", text="时间")
        captures_tree.heading("description", text="描述")
        captures_tree.heading("action", text="操作")
        
        captures_tree.column("id", width=50, anchor=tk.CENTER)
        captures_tree.column("time", width=150)
        captures_tree.column("description", width=300)
        captures_tree.column("action", width=150)
        
        # 添加数据
        for capture in session["captures"]:
            captures_tree.insert("", tk.END, values=(
                capture["id"],
                capture["time"],
                capture["description"],
                "编辑 | 删除"
            ), tags=(capture["id"],))
        
        captures_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 滚动条
        scrollbar = ttk.Scrollbar(captures_frame, orient=tk.VERTICAL, command=captures_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        captures_tree.configure(yscrollcommand=scrollbar.set)
        
        # 编辑和删除功能
        def edit_capture(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region == "cell":
                row = captures_tree.identify_row(event.y)
                col = captures_tree.identify_column(event.x)
                if col == "#4":  # 操作列
                    x, y, width, height = captures_tree.bbox(row, col)
                    if x < event.x < x + width/2:  # 编辑
                        capture_id = int(captures_tree.item(row)["tags"][0])
                        for capture in session["captures"]:
                            if capture["id"] == capture_id:
                                new_desc = simpledialog.askstring(
                                    "编辑描述", 
                                    "请输入描述:", 
                                    initialvalue=capture["description"]
                                )
                                if new_desc is not None:  # 如果不是取消
                                    capture["description"] = new_desc
                                    # 更新树视图
                                    captures_tree.item(row, values=(
                                        capture["id"],
                                        capture["time"],
                                        capture["description"],
                                        "编辑 | 删除"
                                    ))
                                    # 保存会话
                                    self.save_session(session)
                                break
                    else:  # 删除
                        if messagebox.askyesno("确认删除", "确定要删除这条记录吗?"):
                            capture_id = int(captures_tree.item(row)["tags"][0])
                            # 从列表中删除
                            for i, capture in enumerate(session["captures"]):
                                if capture["id"] == capture_id:
                                    # 删除图片文件
                                    if os.path.exists(capture["image_path"]):
                                        os.remove(capture["image_path"])
                                    # 从数组中移除
                                    del session["captures"][i]
                                    break
                            
                            # 重新编号
                            for i, capture in enumerate(session["captures"]):
                                capture["id"] = i + 1
                            
                            # 清空并重新加载树视图
                            for item in captures_tree.get_children():
                                captures_tree.delete(item)
                            for capture in session["captures"]:
                                captures_tree.insert("", tk.END, values=(
                                    capture["id"],
                                    capture["time"],
                                    capture["description"],
                                    "编辑 | 删除"
                                ), tags=(capture["id"],))
                            
                            # 保存会话
                            self.save_session(session)
        
        captures_tree.bind("<Button-1>", edit_capture)
        
        # 查看截图
        def view_screenshot(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region == "cell" and captures_tree.identify_column(event.x) != "#4":
                row = captures_tree.identify_row(event.y)
                capture_id = int(captures_tree.item(row)["tags"][0])
                for capture in session["captures"]:
                    if capture["id"] == capture_id:
                        # 显示截图
                        img_window = tk.Toplevel(editor_window)
                        img_window.title(f"截图 {capture['id']}")
                        
                        img = Image.open(capture["image_path"])
                        # 缩放图片以适应窗口
                        max_width, max_height = 1024, 768
                        img.thumbnail((max_width, max_height))
                        
                        photo = ImageTk.PhotoImage(img)
                        label = ttk.Label(img_window, image=photo)
                        label.image = photo  # 保持引用
                        label.pack()
                        
                        # 显示描述
                        ttk.Label(img_window, text=capture["description"], font=("SimHei", 10)).pack(pady=5)
                        break
        
        captures_tree.bind("<Double-1>", view_screenshot)
        
        # 按钮区域
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        
        # 生成报告按钮
        def generate_report():
            if not session["captures"]:
                messagebox.showwarning("警告", "没有捕捉记录，无法生成报告")
                return
                
            # 选择格式
            format_window = tk.Toplevel(editor_window)
            format_window.title("选择报告格式")
            format_window.geometry("300x200")
            format_window.resizable(False, False)
            format_window.transient(editor_window)
            format_window.grab_set()
            
            ttk.Label(format_window, text="请选择报告格式:", font=("SimHei", 12)).pack(pady=20)
            
            format_var = tk.StringVar(value="docx")
            
            ttk.Radiobutton(format_window, text="Word (.docx)", variable=format_var, value="docx").pack(anchor=tk.W, padx=50, pady=5)
            ttk.Radiobutton(format_window, text="PDF (.pdf)", variable=format_var, value="pdf").pack(anchor=tk.W, padx=50, pady=5)
            ttk.Radiobutton(format_window, text="Markdown (.md)", variable=format_var, value="md").pack(anchor=tk.W, padx=50, pady=5)
            
            def confirm_format():
                report_format = format_var.get()
                format_window.destroy()
                
                # 选择保存路径
                filename = f"{session['name']}_报告.{report_format}"
                save_path = filedialog.asksaveasfilename(
                    defaultextension=f".{report_format}",
                    filetypes=[(f"{report_format.upper()} 文件", f"*.{report_format}")],
                    initialfile=filename
                )
                
                if save_path:
                    if report_format == "docx":
                        self.generate_docx_report(session, save_path)
                    elif report_format == "pdf":
                        self.generate_pdf_report(session, save_path)
                    elif report_format == "md":
                        self.generate_md_report(session, save_path)
                    
                    if messagebox.askyesno("生成成功", f"报告已生成到:\n{save_path}\n是否打开文件?"):
                        os.startfile(save_path)
            
            ttk.Button(format_window, text="确定", command=confirm_format).pack(pady=10)
        
        ttk.Button(btn_frame, text="生成报告", command=generate_report).pack(side=tk.LEFT, padx=5)
        
        # 作废按钮
        def discard_session():
            if messagebox.askyesno("确认作废", f"确定要作废 '{session['name']}' 吗?\n此操作不可恢复!"):
                # 删除文件
                session_path = os.path.join(self.sessions_dir, f"{session['id']}.json")
                if os.path.exists(session_path):
                    os.remove(session_path)
                
                # 删除截图
                images_dir = os.path.join(self.sessions_dir, session["id"])
                if os.path.exists(images_dir):
                    for img_file in os.listdir(images_dir):
                        os.remove(os.path.join(images_dir, img_file))
                    os.rmdir(images_dir)
                
                # 更新列表并关闭窗口
                self.history_sessions.remove(session)
                self.update_history_list()
                editor_window.destroy()
        
        ttk.Button(btn_frame, text="作废", command=discard_session).pack(side=tk.LEFT, padx=5)
    
    def generate_docx_report(self, session, save_path):
        doc = Document()
        
        # 标题
        doc.add_heading("操作记录报告", 0)
        
        # 基本信息
        doc.add_heading("一、基本信息", level=1)
        
        info_table = doc.add_table(rows=5, cols=2)
        info_table.cell(0, 0).text = "操作名称"
        info_table.cell(0, 1).text = session["name"]
        info_table.cell(1, 0).text = "操作描述"
        info_table.cell(1, 1).text = session["description"]
        info_table.cell(2, 0).text = "开始时间"
        info_table.cell(2, 1).text = session["start_time"]
        info_table.cell(3, 0).text = "结束时间"
        info_table.cell(3, 1).text = session["end_time"]
        info_table.cell(4, 0).text = "操作时长"
        info_table.cell(4, 1).text = self.format_duration(session["duration"])
        
        # 操作步骤
        doc.add_heading("二、操作步骤", level=1)
        doc.add_paragraph(f"共 {len(session['captures'])} 步操作")
        
        for i, capture in enumerate(session["captures"], 1):
            doc.add_heading(f"步骤 {i}", level=2)
            doc.add_paragraph(capture["description"])
            
            # 添加图片
            try:
                doc.add_picture(capture["image_path"], width=Inches(6))
                doc.add_caption(f"图 {i}: {capture['description']}")
            except Exception as e:
                doc.add_paragraph(f"[无法加载图片: {str(e)}]")
            
            doc.add_page_break()
        
        # 保存文档
        doc.save(save_path)
    
    def generate_pdf_report(self, session, save_path):
        doc = SimpleDocTemplate(save_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # 标题
        elements.append(Paragraph("操作记录报告", styles['Title']))
        elements.append(Spacer(1, 12))
        
        # 基本信息
        elements.append(Paragraph("一、基本信息", styles['Heading1']))
        elements.append(Spacer(1, 12))
        
        elements.append(Paragraph(f"操作名称: {session['name']}", styles['Normal']))
        elements.append(Paragraph(f"操作描述: {session['description']}", styles['Normal']))
        elements.append(Paragraph(f"开始时间: {session['start_time']}", styles['Normal']))
        elements.append(Paragraph(f"结束时间: {session['end_time']}", styles['Normal']))
        elements.append(Paragraph(f"操作时长: {self.format_duration(session['duration'])}", styles['Normal']))
        elements.append(Spacer(1, 12))
        
        # 操作步骤
        elements.append(Paragraph("二、操作步骤", styles['Heading1']))
        elements.append(Paragraph(f"共 {len(session['captures'])} 步操作", styles['Normal']))
        elements.append(Spacer(1, 12))
        
        for i, capture in enumerate(session["captures"], 1):
            elements.append(Paragraph(f"步骤 {i}", styles['Heading2']))
            elements.append(Paragraph(capture["description"], styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # 添加图片
            try:
                img = RLImage(capture["image_path"], width=400, height=300)
                elements.append(img)
                elements.append(Paragraph(f"图 {i}: {capture['description']}", styles['Italic']))
            except Exception as e:
                elements.append(Paragraph(f"[无法加载图片: {str(e)}]", styles['Normal']))
            
            elements.append(Spacer(1, 24))
        
        # 生成PDF
        doc.build(elements)
    
    def generate_md_report(self, session, save_path):
        with open(save_path, "w", encoding="utf-8") as f:
            # 标题
            f.write("# 操作记录报告\n\n")
            
            # 基本信息
            f.write("## 一、基本信息\n\n")
            f.write(f"| 项目 | 内容 |\n")
            f.write(f"|------|------|\n")
            f.write(f"| 操作名称 | {session['name']} |\n")
            f.write(f"| 操作描述 | {session['description']} |\n")
            f.write(f"| 开始时间 | {session['start_time']} |\n")
            f.write(f"| 结束时间 | {session['end_time']} |\n")
            f.write(f"| 操作时长 | {self.format_duration(session['duration'])} |\n\n")
            
            # 操作步骤
            f.write("## 二、操作步骤\n\n")
            f.write(f"共 {len(session['captures'])} 步操作\n\n")
            
            for i, capture in enumerate(session["captures"], 1):
                f.write(f"### 步骤 {i}\n\n")
                f.write(f"{capture['description']}\n\n")
                f.write(f"![图 {i}: {capture['description']}]({capture['image_path']})\n\n")
                f.write("---\n\n")

if __name__ == "__main__":
    root = tk.Tk()
    app = ScreenCaptureTool(root)
    root.mainloop()
