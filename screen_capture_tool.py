#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
屏幕捕捉工具 - Screen Capture Tool

版权所有 (C) 2025 Unicorn8521
https://github.com/Unicorn8521/screen-capture-tool

本程序是一个功能强大的屏幕捕捉工具，支持以下特性：
- 自定义快捷键截图
- 截图编辑与标注
- 会话管理与历史记录
- 多种格式报告生成（PDF、Word、Markdown）

本程序免费开源，欢迎使用和改进。
"""

import datetime
import json
import os
import sys
import time
import tkinter as tk
import uuid
from tkinter import font as tkfont
from tkinter import ttk, filedialog, messagebox, simpledialog, scrolledtext

import keyboard
from PIL import Image, ImageTk, ImageGrab, ImageOps, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image as RLImage, Spacer


# 常量定义 - 集中管理配置参数
class Config:
    APP_TITLE = "屏幕捕捉工具"
    DEFAULT_WINDOW_SIZE = "600x500"
    EDITOR_WINDOW_SIZE = "900x700"
    PREVIEW_WINDOW_SIZE = "800x600"
    DEFAULT_HOTKEY = "ctrl+alt+o"
    SESSIONS_DIR = os.path.join(os.path.expanduser("~"), ".screen_capture_sessions")
    CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".screen_capture_config.json")
    FONT_FAMILIES = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS", "Microsoft YaHei", "Arial",
                     "sans-serif"]
    TEMP_TIP_DURATION = 2000  # 临时提示显示时间(ms)
    REPORT_FORMATS = {
        "docx": {"desc": "兼容性好，支持编辑", "title": "Word 文档 (.docx)"},
        "pdf": {"desc": "格式固定，跨平台", "title": "PDF 文档 (.pdf)"},
        "md": {"desc": "轻量文本，支持代码块", "title": "Markdown 文档 (.md)"}
    }


# 工具类 - 封装通用功能
class Utils:
    @staticmethod
    def format_duration(seconds):
        """格式化时长为易读字符串"""
        minutes, seconds = divmod(seconds, 60)
        hours, minutes = divmod(minutes, 60)
        if hours > 0:
            return f"{hours}h{minutes}m"
        elif minutes > 0:
            return f"{minutes}m{seconds}s"
        else:
            return f"{seconds}s"

    @staticmethod
    def get_timestamp():
        """获取当前时间戳字符串"""
        return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    @staticmethod
    def get_file_timestamp():
        """获取用于文件名的时间戳"""
        return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    @staticmethod
    def ensure_dir(path):
        """确保目录存在，不存在则创建"""
        if not os.path.exists(path):
            try:
                os.makedirs(path)
                return True
            except Exception as e:
                messagebox.showerror("目录错误", f"创建目录失败：{str(e)}")
                return False
        return True

    @staticmethod
    def get_relative_path(base_path, target_path):
        """计算目标路径相对于基准路径的相对路径"""
        try:
            return os.path.relpath(target_path, base_path)
        except ValueError:
            return target_path  # 跨盘符时返回绝对路径

    @staticmethod
    def get_font_with_chinese_support(size):
        """获取支持中文的字体，按优先级尝试"""
        # 字体列表，包括字体名称和可能的系统路径
        fonts_to_try = [
            # 优先尝试系统字体名称
            "SimHei",
            "WenQuanYi Micro Hei",
            "Heiti TC",
            "Arial Unicode MS",
            "Microsoft YaHei",
            "Arial",
            # 然后尝试具体的文件路径
            "C:/Windows/Fonts/simhei.ttf",  # Windows黑体
            "C:/Windows/Fonts/msyh.ttf",  # Windows微软雅黑
            "C:/Windows/Fonts/simsun.ttc",  # Windows宋体
            "/Library/Fonts/SimHei.ttf",  # macOS黑体
            "/Library/Fonts/Microsoft YaHei.ttc",  # macOS微软雅黑
            "/usr/share/fonts/opentype/noto/NotoSansSC-Regular.ttf",  # Linux Noto Sans
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # Linux文泉驿正黑
        ]

        for font_source in fonts_to_try:
            try:
                # 尝试加载字体
                if os.path.exists(font_source):
                    # 如果是文件路径
                    font = ImageFont.truetype(font_source, size, encoding="utf-8")
                else:
                    # 如果是字体名称
                    font = ImageFont.truetype(font_source, size, encoding="utf-8")
                print(f"成功加载支持中文的字体: {font_source}, 大小: {size}")
                return font
            except (OSError, IOError):
                continue

        # 如果所有字体都加载失败，回退到默认字体
        print(f"所有指定字体都加载失败，使用默认字体，大小: {size}")
        return ImageFont.load_default()


# 报告生成器 - 按格式拆分，单一职责
class DocxReportGenerator:
    @staticmethod
    def generate(session, save_path):
        doc = Document()
        # 报告标题
        doc.add_heading("操作记录报告", 0)

        # 基本信息
        doc.add_heading("一、基本信息", level=1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.cell(0, 0).text = "操作名称"
        info_table.cell(0, 1).text = session["name"]
        info_table.cell(1, 0).text = "操作描述"
        info_table.cell(1, 1).text = session["description"]
        info_table.cell(2, 0).text = "开始时间"
        info_table.cell(2, 1).text = session["start_time"]
        info_table.cell(3, 0).text = "结束时间"
        info_table.cell(3, 1).text = session["end_time"]
        info_table.cell(4, 0).text = "操作时长"
        info_table.cell(4, 1).text = Utils.format_duration(session["duration"])

        # 操作步骤
        doc.add_heading("二、操作步骤", level=1)
        doc.add_paragraph(f"共 {len(session['captures'])} 步操作")
        doc.add_paragraph("")  # 空行分隔

        for i, capture in enumerate(session["captures"], 1):
            # 步骤标题
            step_title = doc.add_paragraph()
            step_title_run = step_title.add_run(f"步骤 {i}: {capture['description']}")
            step_title_run.bold = False
            step_title_run.font.size = Pt(11)
            step_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, capture in enumerate(session["captures"], 1):
            # 步骤标题
            step_title = doc.add_paragraph()
            step_title_run = step_title.add_run(f"步骤 {i}")
            step_title_run.bold = True
            step_title_run.font.size = Pt(14)
            step_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 步骤描述
            doc.add_paragraph(capture["description"])

            try:
                if os.path.exists(capture["image_path"]):
                    # 添加图片
                    doc.add_picture(capture["image_path"], width=Inches(6))
                    # 图片标题
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(f"图 {i}: {capture['description']}")
                    run.font.size = Pt(11)
            except Exception as e:
                doc.add_paragraph(f"[加载图片失败: {str(e)}]")

            # 步骤分页
            # doc.add_page_break()

        doc.save(save_path)
        return True


class PdfReportGenerator:
    @staticmethod
    def _get_available_font():
        """获取可用的中文字体路径"""
        font_candidates = []
        if sys.platform == "win32":
            font_candidates = [
                "C:/Windows/Fonts/simsun.ttc",  # 宋体
                "C:/Windows/Fonts/simhei.ttf",  # 黑体
                "C:/Windows/Fonts/microsoftyahei.ttf"  # 微软雅黑
            ]
        elif sys.platform == "darwin":
            font_candidates = [
                "/Library/Fonts/Songti.ttc",  # 宋体
                "/Library/Fonts/Heiti TC.ttc",  # 黑体
                "/Library/Fonts/Microsoft YaHei.ttc"  # 微软雅黑
            ]
        elif sys.platform == "linux":
            font_candidates = [
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",  # 文泉驿正黑
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",  # Noto Sans
                "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"  # Droid Sans
            ]

        for candidate in font_candidates:
            if os.path.exists(candidate):
                return candidate

        raise FileNotFoundError(
            f"未找到可用中文字体文件，请安装字体后重试。\n搜索路径：\n" + "\n".join(font_candidates)
        )

    @staticmethod
    def generate(session, save_path):
        font_path = PdfReportGenerator._get_available_font()
        pdfmetrics.registerFont(TTFont('Chinese', font_path))
        pdfmetrics.registerFontFamily('Chinese', normal='Chinese')

        doc = SimpleDocTemplate(
            save_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        styles = getSampleStyleSheet()

        # 自定义中文样式
        styles.add(ParagraphStyle(
            name='ChineseTitle',
            fontName='Chinese',
            fontSize=24,
            spaceAfter=20,
            alignment=1
        ))
        styles.add(ParagraphStyle(
            name='ChineseHeading1',
            fontName='Chinese',
            fontSize=16,
            spaceAfter=12,
            spaceBefore=12,
            bold=True
        ))
        styles.add(ParagraphStyle(
            name='ChineseHeading2',
            fontName='Chinese',
            fontSize=14,
            spaceAfter=8,
            spaceBefore=8,
            bold=True
        ))
        styles.add(ParagraphStyle(
            name='ChineseNormal',
            fontName='Chinese',
            fontSize=12,
            spaceAfter=6,
            leading=18
        ))
        styles.add(ParagraphStyle(
            name='ChineseCaption',
            fontName='Chinese',
            fontSize=10,
            spaceAfter=12,
            italic=True,
            alignment=1
        ))

        # 构建内容
        elements = []
        elements.append(Paragraph("操作记录报告", styles['ChineseTitle']))
        elements.append(Paragraph(
            f"报告生成时间：{Utils.get_timestamp()}",
            styles['ChineseNormal']
        ))
        elements.append(Spacer(1, 12))

        # 基本信息
        elements.append(Paragraph("一、操作基本信息", styles['ChineseHeading1']))
        elements.append(Spacer(1, 8))
        info_items = [
            f"操作名称：{session['name']}",
            f"操作描述：{session['description'] or '无'}",
            f"开始时间：{session['start_time']}",
            f"结束时间：{session['end_time']}",
            f"操作时长：{Utils.format_duration(session['duration'])}"
        ]
        for item in info_items:
            elements.append(Paragraph(item, styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        # 操作步骤
        elements.append(Paragraph("二、详细操作步骤", styles['ChineseHeading1']))
        elements.append(Paragraph(f"共 {len(session['captures'])} 步操作", styles['ChineseNormal']))
        elements.append(Spacer(1, 12))

        for i, capture in enumerate(session["captures"], 1):
            elements.append(Paragraph(f"步骤 {i}", styles['ChineseHeading2']))
            elements.append(Paragraph(f"描述：{capture['description']}", styles['ChineseNormal']))
            elements.append(Paragraph(f"截图时间：{capture['time']}", styles['ChineseNormal']))
            elements.append(Spacer(1, 8))

            # 插入截图
            try:
                if os.path.exists(capture["image_path"]):
                    with Image.open(capture["image_path"]) as img:  # 使用with确保资源释放
                        img_width, img_height = img.size
                        scale = min(400 / img_width, 1.0)
                        new_width = int(img_width * scale)
                        new_height = int(img_height * scale)
                        pdf_img = RLImage(capture["image_path"], width=new_width, height=new_height)
                        elements.append(pdf_img)
                        elements.append(Paragraph(f"图 {i}：步骤 {i} 截图", styles['ChineseCaption']))
                else:
                    elements.append(Paragraph(f"[截图文件不存在：{capture['image_path']}]", styles['ChineseNormal']))
            except Exception as e:
                elements.append(Paragraph(f"[截图加载失败：{str(e)}]", styles['ChineseNormal']))

            if i != len(session["captures"]):
                elements.append(Spacer(1, 24))
                elements.append(Paragraph("-" * 60, styles['ChineseNormal']))
                elements.append(Spacer(1, 24))

        doc.build(elements)


class MdReportGenerator:
    @staticmethod
    def generate(session, save_path, use_relative=False):
        """生成Markdown格式报告"""
        report_dir = os.path.dirname(save_path)
        img_rel_dir = "images"
        img_target_dir = os.path.join(report_dir, img_rel_dir)

        # 如果使用相对路径，复制图片到目标目录
        if use_relative and not Utils.ensure_dir(img_target_dir):
            use_relative = False  # 目录创建失败则使用绝对路径

        # 构建Markdown内容
        md_content = []
        md_content.append(f"# 操作记录报告")
        md_content.append(f"**报告生成时间**：{Utils.get_timestamp()}\n")

        # 基本信息
        md_content.append("## 一、操作基本信息")
        md_content.append(f"- **操作名称**：{session['name']}")
        md_content.append(f"- **操作描述**：{session['description'] or '无'}")
        md_content.append(f"- **开始时间**：{session['start_time']}")
        md_content.append(f"- **结束时间**：{session['end_time']}")
        md_content.append(f"- **操作时长**：{Utils.format_duration(session['duration'])}\n")

        # 操作步骤
        md_content.append("## 二、详细操作步骤")
        md_content.append(f"共 {len(session['captures'])} 步操作\n")

        for i, capture in enumerate(session["captures"], 1):
            md_content.append(f"### 步骤 {i}")
            md_content.append(f"- **描述**：{capture['description']}")
            md_content.append(f"- **截图时间**：{capture['time']}")

            try:
                if os.path.exists(capture["image_path"]):
                    if use_relative:
                        # 复制图片到相对路径目录
                        img_filename = os.path.basename(capture["image_path"])
                        target_path = os.path.join(img_target_dir, img_filename)

                        # 仅在文件不存在时复制
                        if not os.path.exists(target_path):
                            with Image.open(capture["image_path"]) as img:
                                img.save(target_path)
                        img_path = f"{img_rel_dir}/{img_filename}"
                    else:
                        img_path = capture["image_path"]

                    md_content.append(f"![图 {i}：步骤 {i} 截图]({img_path})\n")
                else:
                    md_content.append(f"[截图文件不存在：{capture['image_path']}]\n")
            except Exception as e:
                md_content.append(f"[截图加载失败：{str(e)}]\n")

        # 写入文件
        with open(save_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_content))


class ScreenCaptureTool:
    def __init__(self, root):
        self.root = root
        self.root.title(f"{Config.APP_TITLE} - 就绪")
        self.root.geometry(Config.DEFAULT_WINDOW_SIZE)
        self.root.resizable(True, True)

        # 数据存储
        self.current_session_id = None
        self.current_session = self._init_empty_session()
        self.history_sessions = []
        self.sessions_dir = Config.SESSIONS_DIR

        # 状态变量
        self.is_capturing = False
        self.hotkey = Config.DEFAULT_HOTKEY
        self.hotkey_obj = None
        self.start_time = 0
        self.images_dir = ""

        # 初始化字体
        self._init_fonts()

        # 确保会话目录存在
        Utils.ensure_dir(self.sessions_dir)

        # 加载快捷键配置
        self.load_hotkey_config()

        # 加载历史记录
        self.load_history_sessions()

        # 绑定关闭事件
        self.root.protocol("WM_DELETE_WINDOW", self.on_main_window_close)

        # 创建界面
        self.create_main_interface()

    def _init_empty_session(self):
        """初始化空会话数据结构"""
        return {
            "id": "",
            "name": "",
            "description": "",
            "start_time": "",
            "end_time": "",
            "duration": 0,
            "captures": []
        }

    def _init_fonts(self):
        """初始化应用字体"""
        default_font = tkfont.nametofont("TkDefaultFont")
        default_font.configure(family=Config.FONT_FAMILIES)
        self.root.option_add("*Font", default_font)

    def create_main_interface(self):
        """创建主界面"""
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题与状态
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=5)
        ttk.Label(title_frame, text=Config.APP_TITLE, font=(Config.FONT_FAMILIES[0], 16)).pack(side=tk.LEFT)
        self.status_label = ttk.Label(
            title_frame,
            text="状态：就绪",
            font=(Config.FONT_FAMILIES[0], 10),
            foreground="#666"
        )
        self.status_label.pack(side=tk.RIGHT)

        # 控制按钮
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(pady=10)
        self.start_btn = ttk.Button(control_frame, text="开始捕捉", command=self.start_capture)
        self.start_btn.pack(side=tk.LEFT, padx=5)
        self.stop_btn = ttk.Button(control_frame, text="停止捕捉", command=self.stop_capture, state=tk.DISABLED)
        self.stop_btn.pack(side=tk.LEFT, padx=5)

        # 快捷键设置区域
        hotkey_frame = ttk.Frame(main_frame)
        hotkey_frame.pack(pady=5, fill=tk.X)
        ttk.Label(hotkey_frame, text="截图快捷键：", font=(Config.FONT_FAMILIES[0], 10)).pack(side=tk.LEFT, padx=5)
        self.hotkey_label = ttk.Label(hotkey_frame, text=self.hotkey.upper(),
                                      font=(Config.FONT_FAMILIES[0], 10, "bold"), foreground="#1976d2")
        self.hotkey_label.pack(side=tk.LEFT, padx=5)
        ttk.Button(hotkey_frame, text="设置快捷键", command=self.set_hotkey).pack(side=tk.LEFT, padx=5)

        # 历史记录区域
        ttk.Label(main_frame, text="历史记录", font=(Config.FONT_FAMILIES[0], 12)).pack(anchor=tk.W, pady=5)
        history_frame = ttk.Frame(main_frame)
        history_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 历史记录表格
        columns = ("name", "date", "duration", "captures")
        self.history_tree = ttk.Treeview(history_frame, columns=columns, show="headings")
        self.history_tree.heading("name", text="操作名称")
        self.history_tree.heading("date", text="开始时间")
        self.history_tree.heading("duration", text="时长")
        self.history_tree.heading("captures", text="截图数")
        self.history_tree.column("name", width=180)
        self.history_tree.column("date", width=180)
        self.history_tree.column("duration", width=80, anchor=tk.CENTER)
        self.history_tree.column("captures", width=80, anchor=tk.CENTER)
        self.history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 滚动条
        scrollbar = ttk.Scrollbar(history_frame, orient=tk.VERTICAL, command=self.history_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.history_tree.configure(yscrollcommand=scrollbar.set)

        # 历史记录右键菜单
        self.history_menu = tk.Menu(self.root, tearoff=0)
        self.history_menu.add_command(label="重命名", command=self.rename_session)
        self.history_menu.add_command(label="删除", command=self.delete_session)

        # 绑定事件
        self.history_tree.bind("<Double-1>", self.open_session)
        self.history_tree.bind("<Button-3>", self.show_history_menu)

        # 更新历史记录列表
        self.update_history_list()

        # 版权标识
        footer_frame = ttk.Frame(main_frame)
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=5)
        ttk.Label(
            footer_frame,
            text="© 2025 Unicorn8521 版权所有",
            font=(Config.FONT_FAMILIES[0], 9),
            foreground="#999"
        ).pack(side=tk.RIGHT, padx=10)

    def show_history_menu(self, event):
        """显示历史记录右键菜单"""
        item = self.history_tree.identify_row(event.y)
        if item:
            self.history_tree.selection_set(item)
            self.history_menu.post(event.x_root, event.y_root)

    def rename_session(self):
        """重命名会话"""
        selected_item = self.history_tree.selection()
        if not selected_item:
            messagebox.showwarning("提示", "请先选择一条历史记录")
            return

        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                new_name = simpledialog.askstring(
                    "重命名",
                    "请输入新操作名称：",
                    initialvalue=session["name"]
                )
                if new_name and new_name.strip():
                    session["name"] = new_name.strip()
                    self.save_session(session)
                    self.update_history_list()
                    messagebox.showinfo("成功", "操作名称已更新")
                break

    def delete_session(self):
        """删除会话"""
        selected_item = self.history_tree.selection()
        if not selected_item:
            messagebox.showwarning("提示", "请先选择一条历史记录")
            return

        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                if messagebox.askyesno(
                        "确认删除",
                        f"确定删除「{session['name']}」及所有截图？\n此操作不可恢复！"
                ):
                    self.delete_session_by_id(session_id)
                break

    def delete_session_by_id(self, session_id):
        """通过ID删除会话"""
        # 查找会话
        session_to_remove = None
        for session in self.history_sessions:
            if session["id"] == session_id:
                session_to_remove = session
                break

        if not session_to_remove:
            return

        # 删除JSON配置文件
        session_path = os.path.join(self.sessions_dir, f"{session_id}.json")
        if os.path.exists(session_path):
            try:
                os.remove(session_path)
            except Exception as e:
                messagebox.showerror("错误", f"删除配置文件失败：{str(e)}")
                return

        # 删除截图文件夹
        images_dir = os.path.join(self.sessions_dir, session_id)
        if os.path.exists(images_dir):
            try:
                for img_file in os.listdir(images_dir):
                    os.remove(os.path.join(images_dir, img_file))
                os.rmdir(images_dir)
            except Exception as e:
                messagebox.showerror("错误", f"删除截图文件失败：{str(e)}")
                return

        # 更新列表
        self.history_sessions.remove(session_to_remove)
        self.update_history_list()
        messagebox.showinfo("成功", "历史记录已删除")

    def start_capture(self):
        """开始捕捉会话"""
        if self.is_capturing:
            messagebox.showwarning("提示", "捕捉已在运行中，无需重复启动")
            return

        # 获取会话名称
        session_name = simpledialog.askstring("操作名称", "请输入此次捕捉的操作名称：")
        if not session_name or not session_name.strip():
            messagebox.showwarning("提示", "操作名称不能为空")
            return
        session_name = session_name.strip()

        # 初始化会话数据
        self.current_session_id = str(uuid.uuid4())
        self.current_session = {
            "id": self.current_session_id,
            "name": session_name,
            "description": simpledialog.askstring("操作描述", "请输入此次捕捉的描述（可选）：") or "",
            "start_time": Utils.get_timestamp(),
            "end_time": "",
            "duration": 0,
            "captures": []
        }

        # 创建截图存储目录
        self.images_dir = os.path.join(self.sessions_dir, self.current_session_id)
        if not Utils.ensure_dir(self.images_dir):
            return

        # 注册快捷键
        try:
            self.hotkey_obj = keyboard.add_hotkey(self.hotkey, self.capture_screen)
        except PermissionError:
            messagebox.showerror(
                "权限错误",
                "快捷键注册失败！\n请以管理员身份运行程序（Windows）或授予键盘访问权限（macOS）"
            )
            return
        except Exception as e:
            messagebox.showerror("错误", f"快捷键注册失败：{str(e)}")
            return

        # 更新状态
        self.is_capturing = True
        self.start_time = time.time()
        self.start_btn.config(state=tk.DISABLED)
        self.stop_btn.config(state=tk.NORMAL)
        self.root.title(f"{Config.APP_TITLE} - 捕捉中（{self.hotkey.upper()}截图）")
        self.status_label.config(
            text=f"状态：捕捉中（{self.hotkey.upper()}截图）",
            foreground="#d32f2f"
        )

        # 最小化主窗口并提示
        self.root.iconify()
        messagebox.showinfo(
            "捕捉启动",
            f"捕捉已开始！\n• 按 {self.hotkey.upper()} 截取当前屏幕\n• 点击「停止捕捉」结束操作\n• 截图自动保存至：{self.images_dir}"
        )

    def stop_capture(self):
        """停止捕捉会话"""
        if not self.is_capturing:
            return

        # 清理快捷键
        if self.hotkey_obj:
            try:
                keyboard.remove_hotkey(self.hotkey_obj)
                self.hotkey_obj = None
            except Exception as e:
                messagebox.showwarning("警告", f"快捷键清理失败：{str(e)}")

        # 更新会话数据
        self.is_capturing = False
        self.current_session["end_time"] = Utils.get_timestamp()
        self.current_session["duration"] = int(time.time() - self.start_time)

        # 恢复界面状态
        self.start_btn.config(state=tk.NORMAL)
        self.stop_btn.config(state=tk.DISABLED)
        self.root.title(f"{Config.APP_TITLE} - 就绪")
        self.status_label.config(text="状态：就绪", foreground="#666")

        # 保存会话并刷新历史
        self.save_current_session()
        self.load_history_sessions()
        self.update_history_list()

        # 打开编辑窗口
        if self.current_session["captures"]:
            self.open_editor_window(self.current_session)
        else:
            messagebox.showinfo("提示", "此次捕捉未生成任何截图，无需编辑")

    def capture_screen(self):
        """捕捉屏幕截图"""
        if not self.is_capturing:
            return

        try:
            # 确保目录存在
            if not os.path.exists(self.images_dir):
                os.makedirs(self.images_dir)

            # 捕获全屏截图
            screenshot = ImageGrab.grab()
            original_size = screenshot.size
        except ImportError:
            messagebox.showerror(
                "依赖缺失",
                "Linux 系统需安装 python3-xlib 才能截图\n执行命令：sudo apt install python3-xlib"
            )
            self.stop_capture()
            return
        except Exception as e:
            messagebox.showerror("截图错误", f"截图失败：{str(e)}")
            return

        # 显示预览窗口并让用户选择是否裁剪和添加描述
        result = self.show_capture_preview(screenshot)
        if not result:
            # 用户取消了截图保存
            return

        final_image, description = result
        final_size = final_image.size

        # 调试信息：检查是否进行了裁剪
        is_cropped = final_size != original_size
        print(f"原始图像大小: {original_size}, 保存图像大小: {final_size}, 是否已裁剪: {is_cropped}")

        # 保存截图
        capture_time = Utils.get_file_timestamp()
        img_filename = f"capture_{capture_time}.png"
        img_path = os.path.join(self.images_dir, img_filename)
        try:
            # 确保使用裁剪后的图像保存
            final_image.save(img_path)
            print(f"已保存图像到: {img_path}")

            # 验证保存是否成功
            if os.path.exists(img_path):
                saved_img = Image.open(img_path)
                saved_size = saved_img.size
                print(f"保存的图像实际大小: {saved_size}")
                saved_img.close()
        except Exception as e:
            messagebox.showerror("保存错误", f"截图保存失败：{str(e)}")
            print(f"保存错误详情: {str(e)}")
            return

        # 添加到会话记录
        capture_count = len(self.current_session["captures"]) + 1
        self.current_session["captures"].append({
            "id": capture_count,
            "time": Utils.get_timestamp(),
            "description": description or f"第{capture_count}次记录",
            "image_path": img_path
        })

        # 显示临时提示，包含是否裁剪的信息
        if is_cropped:
            self.show_temp_tip(f"已完成第{capture_count}次截图 (已裁剪)")
        else:
            self.show_temp_tip(f"已完成第{capture_count}次截图")

    def show_capture_preview(self, screenshot, description=""):
        """显示截图预览窗口，支持区域选择、添加描述、自由标注和鼠标滚轮缩放功能"""
        # 创建全屏预览窗口
        preview_window = tk.Toplevel(self.root)
        preview_window.attributes("-fullscreen", True)
        preview_window.attributes("-topmost", True)

        # 创建半透明背景画布
        canvas = tk.Canvas(preview_window, cursor="cross")
        canvas.pack(fill=tk.BOTH, expand=True)

        # 将截图转换为PhotoImage
        screen_width = preview_window.winfo_screenwidth()
        screen_height = preview_window.winfo_screenheight()

        # 调整图像大小以适应屏幕
        img_width, img_height = screenshot.size
        scale = min(screen_width / img_width, screen_height / img_height, 1.0)
        new_width = int(img_width * scale)
        new_height = int(img_height * scale)

        resized_img = screenshot.resize((new_width, new_height), Image.Resampling.LANCZOS)
        photo = ImageTk.PhotoImage(resized_img)

        # 居中显示图像
        x_pos = (screen_width - new_width) // 2
        y_pos = (screen_height - new_height) // 2
        image_id = canvas.create_image(x_pos, y_pos, anchor=tk.NW, image=photo)

        # 存储原始图像和缩放信息
        canvas.original_image = screenshot
        canvas.original_width = img_width  # 保存原始图像宽度
        canvas.original_height = img_height  # 保存原始图像高度
        canvas.resized_image = resized_img
        canvas.photo = photo
        canvas.scale = scale  # 当前缩放比例
        canvas.x_pos = x_pos  # 图像在画布上的X坐标
        canvas.y_pos = y_pos  # 图像在画布上的Y坐标
        canvas.image_id = image_id

        # 滚轮缩放函数
        def on_mousewheel(event):
            # 获取鼠标在画布上的位置
            mouse_x = canvas.canvasx(event.x)
            mouse_y = canvas.canvasy(event.y)

            # 计算鼠标在图像上的相对位置（如果鼠标在图像区域内）
            rel_x = mouse_x - canvas.x_pos
            rel_y = mouse_y - canvas.y_pos

            # 检查鼠标是否在图像区域内
            if 0 <= rel_x < canvas.resized_image.width and 0 <= rel_y < canvas.resized_image.height:
                # 判断缩放方向
                zoom_in = False
                if hasattr(event, 'delta'):
                    zoom_in = event.delta > 0
                else:
                    # Linux系统
                    zoom_in = event.num == 4

                if zoom_in:
                    new_scale = min(canvas.scale * 1.1, 5.0)  # 最大放大5倍
                else:
                    new_scale = max(canvas.scale / 1.1, 0.1)  # 最小缩小到0.1倍

                # 计算缩放前鼠标在原图上的相对位置
                orig_rel_x = rel_x / canvas.scale
                orig_rel_y = rel_y / canvas.scale

                # 更新缩放比例
                canvas.scale = new_scale

                # 计算新的图像大小
                new_img_width = int(canvas.original_width * new_scale)
                new_img_height = int(canvas.original_height * new_scale)

                # 重新调整图像大小
                canvas.resized_image = canvas.original_image.resize((new_img_width, new_img_height),
                                                                    Image.Resampling.LANCZOS)
                canvas.photo = ImageTk.PhotoImage(canvas.resized_image)

                # 计算新的图像位置，保持鼠标指向的原图位置不变
                new_x_pos = mouse_x - (orig_rel_x * new_scale)
                new_y_pos = mouse_y - (orig_rel_y * new_scale)

                # 更新图像显示
                canvas.coords(canvas.image_id, new_x_pos, new_y_pos)
                canvas.itemconfig(canvas.image_id, image=canvas.photo)

                # 更新位置信息
                canvas.x_pos = new_x_pos
                canvas.y_pos = new_y_pos

                # 重新绘制遮罩
                draw_mask()

        # 绑定滚轮事件
        canvas.bind("<MouseWheel>", on_mousewheel)  # Windows
        canvas.bind("<Button-4>", on_mousewheel)  # Linux
        canvas.bind("<Button-5>", on_mousewheel)  # Linux

        # 工具模式：'select'、'draw' 或 'text'
        current_mode = 'select'  # 默认选择模式
        canvas.current_mode = current_mode

        # 区域选择相关变量
        rect_id = None
        start_x = 0
        start_y = 0
        selected_region = None

        # 自由标注相关变量
        drawing = False
        current_line = []
        current_line_id = None
        annotations = []  # 存储所有标注的ID
        annotation_paths = []  # 存储标注路径信息：[(points, color, width), ...]
        text_annotations = []  # 存储文本标注信息：[(x, y, text, color, font_size), ...]
        annotation_color = "red"  # 默认标注颜色
        annotation_width = 2  # 默认线条宽度
        font_size = 12  # 默认字体大小
        current_text_id = None  # 当前正在编辑的文本ID

        # 描述输入，使用传入的描述作为默认值
        description_var = tk.StringVar(value=description)

        # 绘制半透明遮罩
        def draw_mask():
            # 清除所有非图像、非标注元素
            for item in canvas.find_all():
                if item != canvas.image_id and not item in annotations:
                    canvas.delete(item)

            # 绘制四个遮罩区域
            img_width = canvas.resized_image.width
            img_height = canvas.resized_image.height
            canvas.create_rectangle(0, 0, screen_width, canvas.y_pos, fill="black", stipple="gray50")
            canvas.create_rectangle(0, canvas.y_pos, canvas.x_pos, canvas.y_pos + img_height, fill="black",
                                    stipple="gray50")
            canvas.create_rectangle(canvas.x_pos + img_width, canvas.y_pos, screen_width, canvas.y_pos + img_height,
                                    fill="black",
                                    stipple="gray50")
            canvas.create_rectangle(0, canvas.y_pos + img_height, screen_width, screen_height, fill="black",
                                    stipple="gray50")

        draw_mask()

        # 鼠标按下事件
        def on_mouse_down(event):
            nonlocal rect_id, start_x, start_y, drawing, current_line, current_line_id, current_text_id
            # 检查点击是否在图像区域内 - 使用canvas的当前属性
            if canvas.x_pos <= event.x <= canvas.x_pos + canvas.resized_image.width and \
                    canvas.y_pos <= event.y <= canvas.y_pos + canvas.resized_image.height:
                start_x = event.x
                start_y = event.y

                if current_mode == 'select':
                    # 选择模式：创建选择矩形
                    rect_id = canvas.create_rectangle(start_x, start_y, start_x, start_y, outline="blue", width=2)
                elif current_mode == 'draw':
                    # 绘制模式：开始自由绘制
                    drawing = True
                    current_line = [(start_x, start_y)]
                    # 创建初始点（小圆圈）
                    current_line_id = canvas.create_oval(
                        start_x - 1, start_y - 1, start_x + 1, start_y + 1,
                        fill=annotation_color, outline=annotation_color, width=annotation_width
                    )
                    annotations.append(current_line_id)
                elif current_mode == 'text':
                    # 文本模式：在点击位置添加文本标注
                    # 临时关闭预览窗口的置顶属性，以便文本输入框能够弹出
                    preview_window.attributes("-topmost", False)

                    # 获取用户输入的文本
                    text = simpledialog.askstring("文本标注", "请输入要添加的文本：")

                    # 重新启用预览窗口的置顶属性
                    preview_window.attributes("-topmost", True)

                    if text and text.strip():
                        # 在画布上显示文本
                        text_id = canvas.create_text(start_x, start_y, text=text.strip(),
                                                     fill=annotation_color, font=(Config.FONT_FAMILIES[0], font_size),
                                                     anchor=tk.NW)
                        annotations.append(text_id)

                        # 计算文本在原始图像上的坐标 - 使用当前缩放比例
                        orig_x = int((start_x - canvas.x_pos) / canvas.scale)
                        orig_y = int((start_y - canvas.y_pos) / canvas.scale)

                        # 保存文本标注信息
                        text_annotations.append((orig_x, orig_y, text.strip(), annotation_color, font_size))

        # 鼠标移动事件
        def on_mouse_move(event):
            nonlocal rect_id, drawing, current_line, current_line_id

            if current_mode == 'select' and rect_id:
                # 选择模式：更新矩形大小
                canvas.coords(rect_id, start_x, start_y, event.x, event.y)
            elif current_mode == 'draw' and drawing:
                # 绘制模式：更新自由绘制线条
                # 使用canvas的当前属性进行判断
                if canvas.x_pos <= event.x <= canvas.x_pos + canvas.resized_image.width and \
                        canvas.y_pos <= event.y <= canvas.y_pos + canvas.resized_image.height:
                    # 记录当前点
                    current_line.append((event.x, event.y))
                    # 绘制从最后一个点到当前点的线段
                    last_x, last_y = current_line[-2]
                    new_line_id = canvas.create_line(
                        last_x, last_y, event.x, event.y,
                        fill=annotation_color, width=annotation_width, smooth=True
                    )
                    # 添加到标注列表
                    annotations.append(new_line_id)
                    # 更新当前线条ID
                    current_line_id = new_line_id

        # 鼠标释放事件
        def on_mouse_up(event):
            nonlocal rect_id, selected_region, drawing, current_line

            if current_mode == 'select' and rect_id:
                # 选择模式：处理区域选择
                # 计算实际选择区域
                x1 = min(start_x, event.x)
                y1 = min(start_y, event.y)
                x2 = max(start_x, event.x)
                y2 = max(start_y, event.y)

                # 确保选中区域有效（最小尺寸检查）
                if (x2 - x1) > 5 and (y2 - y1) > 5:
                    # 转换为原图坐标 - 使用canvas的当前属性和缩放比例
                    orig_x1 = max(0, int((x1 - canvas.x_pos) / canvas.scale))
                    orig_y1 = max(0, int((y1 - canvas.y_pos) / canvas.scale))
                    orig_x2 = min(canvas.original_width, int((x2 - canvas.x_pos) / canvas.scale))
                    orig_y2 = min(canvas.original_height, int((y2 - canvas.y_pos) / canvas.scale))

                    if orig_x2 > orig_x1 and orig_y2 > orig_y1:
                        selected_region = (orig_x1, orig_y1, orig_x2, orig_y2)
                        # 在界面上显示选中区域的大小信息
                        region_info = f"已选择区域：{orig_x2 - orig_x1}x{orig_y2 - orig_y1}像素"
                        # 移除之前可能存在的信息标签
                        for item in canvas.find_withtag("region_info"):
                            canvas.delete(item)
                        # 创建新的信息标签
                        canvas.create_text(screen_width // 2, 20, text=region_info, fill="white",
                                           font=(Config.FONT_FAMILIES[0], 12), tag="region_info")

                # 移除矩形
                canvas.delete(rect_id)
                rect_id = None
            elif current_mode == 'draw' and drawing:
                # 绘制模式：结束绘制
                # 将当前绘制的线条坐标点转换为相对于原始图像的坐标，并保存
                if len(current_line) > 1:
                    # 转换坐标到原始图像坐标系 - 使用canvas的当前属性和缩放比例
                    original_points = []
                    for (px, py) in current_line:
                        # 减去图像在画布上的偏移量，再除以缩放比例
                        orig_x = int((px - canvas.x_pos) / canvas.scale)
                        orig_y = int((py - canvas.y_pos) / canvas.scale)
                        original_points.append((orig_x, orig_y))

                    # 保存标注路径信息
                    annotation_paths.append((original_points, annotation_color, annotation_width))

                drawing = False
                current_line = []

        # 确认按钮回调
        def on_confirm():
            # 获取描述文本
            description = description_var.get().strip()

            # 根据是否选择了区域来裁剪图像
            if selected_region:
                # 确保正确裁剪图像
                final_image = canvas.original_image.crop(selected_region)
                # 可选：添加日志以确认裁剪操作
                print(f"正在保存裁剪图像，区域：{selected_region}")
            else:
                final_image = canvas.original_image.copy()
                print("正在保存完整图像，未进行裁剪")

            # 检查是否有标注或文本需要绘制到最终图像上
            if annotation_paths or text_annotations:
                # 创建一个可以绘制的图像副本
                # 直接导入ImageDraw，避免变量未定义错误
                from PIL import ImageDraw
                draw_image = final_image.copy()
                draw = ImageDraw.Draw(draw_image)

                # 根据是否选择了区域进行不同的坐标处理
                if selected_region:
                    orig_x1, orig_y1, orig_x2, orig_y2 = selected_region
                    # 对于每个标注路径，绘制到裁剪后的图像上
                    for points, color, width in annotation_paths:
                        if len(points) > 1:
                            adjusted_points = []
                            # 只处理裁剪区域内的点
                            for px, py in points:
                                if orig_x1 <= px <= orig_x2 and orig_y1 <= py <= orig_y2:
                                    # 转换为裁剪区域的相对坐标
                                    adjusted_points.append((px - orig_x1, py - orig_y1))
                            # 绘制调整后的线条
                            if len(adjusted_points) > 1:
                                for i in range(len(adjusted_points) - 1):
                                    draw.line(
                                        [adjusted_points[i], adjusted_points[i + 1]],
                                        fill=color,
                                        width=width
                                    )

                    # 绘制文本标注（需要调整坐标）
                    for x, y, text, color, size in text_annotations:
                        if orig_x1 <= x <= orig_x2 and orig_y1 <= y <= orig_y2:
                            # 转换为裁剪区域的相对坐标
                            adjusted_x = x - orig_x1
                            adjusted_y = y - orig_y1
                            # 使用增强的字体加载函数，确保中文支持和正确的字体大小
                            font = Utils.get_font_with_chinese_support(size)
                            # 绘制文本，添加日志验证
                            print(
                                f"正在渲染文本: '{text}', 位置: ({adjusted_x}, {adjusted_y}), 颜色: {color}, 字体大小: {size}")
                            draw.text((adjusted_x, adjusted_y), text, fill=color, font=font)
                else:
                    # 在完整图像上绘制所有标注
                    for points, color, width in annotation_paths:
                        if len(points) > 1:
                            for i in range(len(points) - 1):
                                draw.line(
                                    [points[i], points[i + 1]],
                                    fill=color,
                                    width=width
                                )

                    # 在完整图像上绘制所有文本标注
                    for x, y, text, color, size in text_annotations:
                        # 使用增强的字体加载函数，确保中文支持和正确的字体大小
                        font = Utils.get_font_with_chinese_support(size)
                        # 绘制文本，添加日志验证
                        print(f"正在渲染文本: '{text}', 位置: ({x}, {y}), 颜色: {color}, 字体大小: {size}")
                        draw.text((x, y), text, fill=color, font=font)

                print(
                    f"已保存包含{len(annotation_paths)}个标注路径和{len(text_annotations)}个文本标注的图像，字体设置已正确应用")
                final_image = draw_image

            # 设置返回结果
            preview_window.result = (final_image, description)
            # 关闭窗口
            preview_window.destroy()

        # 取消按钮回调
        def on_cancel():
            preview_window.result = None
            preview_window.destroy()

        # 键盘事件处理
        def on_key_press(event):
            # 只保留必要的Enter和Escape键功能，不处理其他快捷键
            if event.keysym == "Return":
                # 支持单独的Enter键保存
                on_confirm()
            elif event.keysym == "Escape":
                on_cancel()
            # 移除其他所有快捷键处理

        # 绑定事件
        canvas.bind("<Button-1>", on_mouse_down)
        canvas.bind("<B1-Motion>", on_mouse_move)
        canvas.bind("<ButtonRelease-1>", on_mouse_up)
        preview_window.bind("<Key>", on_key_press)

        # 创建工具按钮区域
        tools_frame = ttk.Frame(preview_window, style="TFrame")
        tools_frame.place(relx=0.5, rely=0.9, anchor="s", width=screen_width, height=60)
        tools_frame.configure(style="Transparent.TFrame")

        # 创建底部控制区域
        control_frame = ttk.Frame(preview_window, style="TFrame")
        control_frame.place(relx=0.5, rely=1.0, anchor="s", width=screen_width, height=80)
        control_frame.configure(style="Transparent.TFrame")

        # 设置样式
        style = ttk.Style()
        style.configure("Transparent.TFrame", background="#333333", borderwidth=0)
        style.configure("Control.TButton", font=(Config.FONT_FAMILIES[0], 12), padding=10)
        style.configure("Tool.TButton", font=(Config.FONT_FAMILIES[0], 11), padding=8)
        style.configure("ActiveTool.TButton", background="#4a90e2", foreground="white")

        # 切换工具模式的函数
        def switch_mode(mode):
            nonlocal current_mode
            current_mode = mode
            canvas.current_mode = mode

            # 更新按钮状态
            if mode == 'select':
                select_btn.configure(style="ActiveTool.TButton")
                draw_btn.configure(style="Tool.TButton")
                text_btn.configure(style="Tool.TButton")
                canvas.config(cursor="cross")
            elif mode == 'draw':
                select_btn.configure(style="Tool.TButton")
                draw_btn.configure(style="ActiveTool.TButton")
                text_btn.configure(style="Tool.TButton")
                canvas.config(cursor="pencil")
            elif mode == 'text':
                select_btn.configure(style="Tool.TButton")
                draw_btn.configure(style="Tool.TButton")
                text_btn.configure(style="ActiveTool.TButton")
                canvas.config(cursor="tcross")

        # 清除所有标注的函数
        def clear_annotations():
            for item_id in annotations:
                canvas.delete(item_id)
            annotations.clear()
            annotation_paths.clear()
            text_annotations.clear()

        # 颜色选择函数
        def change_color():
            nonlocal annotation_color
            # 预定义的颜色选项
            color_options = ["red", "blue", "green", "black", "yellow", "purple", "orange", "white"]

            # 创建颜色选择对话框
            color_window = tk.Toplevel(preview_window)
            color_window.title("选择颜色")
            color_window.geometry("300x150")
            color_window.transient(preview_window)
            color_window.grab_set()

            # 显示颜色选项
            color_frame = ttk.Frame(color_window)
            color_frame.pack(pady=20)

            for color in color_options:
                # 创建颜色按钮
                color_btn = ttk.Button(color_frame, text="  ", width=3)
                # 设置按钮背景色
                color_btn.configure(style=f"ColorButton.{color}.TButton")
                # 绑定点击事件
                color_btn.configure(command=lambda c=color: select_color(c, color_window))
                color_btn.pack(side=tk.LEFT, padx=5)

            def select_color(color, window):
                nonlocal annotation_color
                annotation_color = color
                # 更新颜色按钮文本以显示当前选择
                if hasattr(color_window, 'current_color'):
                    color_window.current_color.destroy()
                color_window.current_color = ttk.Label(window, text=f"当前颜色: {color}",
                                                       foreground=color, font=(Config.FONT_FAMILIES[0], 12))
                color_window.current_color.pack(pady=10)

            # 应用样式
            style = ttk.Style()
            for color in color_options:
                style.configure(f"ColorButton.{color}.TButton", background=color)

        # 字体大小选择函数
        def change_font_size():
            nonlocal font_size
            # 创建自定义对话框并设置为置顶
            size_window = tk.Toplevel(preview_window)
            size_window.title("字体大小")
            size_window.geometry("300x120")
            size_window.transient(preview_window)
            size_window.grab_set()
            size_window.attributes("-topmost", True)

            # 确保窗口在最上层
            size_window.update_idletasks()
            size_window.lift()

            # 创建标签和输入框
            ttk.Label(size_window, text="请输入字体大小(1-99):", font=(Config.FONT_FAMILIES[0], 11)).pack(pady=10)

            size_var = tk.StringVar(value=str(font_size))
            size_entry = ttk.Entry(size_window, textvariable=size_var, width=20, font=(Config.FONT_FAMILIES[0], 12))
            size_entry.pack(pady=5)
            size_entry.focus_set()

            # 确认按钮的回调函数
            def on_confirm():
                size_str = size_var.get()
                if size_str and size_str.isdigit():
                    size = int(size_str)
                    if 1 <= size <= 99:
                        nonlocal font_size
                        font_size = size
                        size_window.destroy()
                    else:
                        messagebox.showwarning("输入错误", "字体大小必须在1到99之间！")
                else:
                    messagebox.showwarning("输入错误", "请输入有效的数字！")

            # 创建按钮
            btn_frame = ttk.Frame(size_window)
            btn_frame.pack(pady=10)

            ttk.Button(btn_frame, text="确定", command=on_confirm).pack(side=tk.LEFT, padx=10)
            ttk.Button(btn_frame, text="取消", command=size_window.destroy).pack(side=tk.LEFT, padx=10)

            # 绑定Enter键确认
            size_window.bind('<Return>', lambda event: on_confirm())
            size_window.bind('<Escape>', lambda event: size_window.destroy())

        # 添加工具按钮
        tools_inner_frame = ttk.Frame(tools_frame, style="Transparent.TFrame")
        tools_inner_frame.pack(fill=tk.X, padx=20, pady=5)

        ttk.Label(tools_inner_frame, text="工具：", font=(Config.FONT_FAMILIES[0], 11),
                  background="#333333", foreground="white").pack(side=tk.LEFT, padx=(0, 10))

        select_btn = ttk.Button(tools_inner_frame, text="选择区域", command=lambda: switch_mode('select'),
                                style="ActiveTool.TButton")
        select_btn.pack(side=tk.LEFT, padx=5)

        draw_btn = ttk.Button(tools_inner_frame, text="自由标注", command=lambda: switch_mode('draw'),
                              style="Tool.TButton")
        draw_btn.pack(side=tk.LEFT, padx=5)

        text_btn = ttk.Button(tools_inner_frame, text="文本标注", command=lambda: switch_mode('text'),
                              style="Tool.TButton")
        text_btn.pack(side=tk.LEFT, padx=5)

        color_btn = ttk.Button(tools_inner_frame, text="选择颜色", command=change_color,
                               style="Tool.TButton")
        color_btn.pack(side=tk.LEFT, padx=5)

        font_btn = ttk.Button(tools_inner_frame, text="字体大小", command=change_font_size,
                              style="Tool.TButton")
        font_btn.pack(side=tk.LEFT, padx=5)

        clear_btn = ttk.Button(tools_inner_frame, text="清除标注", command=clear_annotations,
                               style="Tool.TButton")
        clear_btn.pack(side=tk.LEFT, padx=5)

        # 添加描述输入框
        ttk.Label(control_frame, text="描述：", font=(Config.FONT_FAMILIES[0], 12), background="#333333",
                  foreground="white").pack(side=tk.LEFT, padx=10, pady=10, anchor="center")
        description_entry = ttk.Entry(control_frame, textvariable=description_var, width=50,
                                      font=(Config.FONT_FAMILIES[0], 12))
        description_entry.pack(side=tk.LEFT, padx=10, pady=10, anchor="center", fill=tk.X, expand=True)
        description_entry.focus_set()

        # 添加按钮
        button_frame = ttk.Frame(control_frame, style="Transparent.TFrame")
        button_frame.pack(side=tk.RIGHT, padx=20, pady=10)

        cancel_btn = ttk.Button(button_frame, text="取消", command=on_cancel, style="Control.TButton")
        cancel_btn.pack(side=tk.RIGHT, padx=10)

        confirm_btn = ttk.Button(button_frame, text="保存", command=on_confirm, style="Control.TButton")
        confirm_btn.pack(side=tk.RIGHT, padx=10)

        # 添加提示标签
        hint_frame = ttk.Frame(control_frame, style="Transparent.TFrame")
        hint_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=20, pady=(0, 5))

        hint_text = "提示：Enter:保存 | ESC:取消"
        ttk.Label(hint_frame, text=hint_text, font=(Config.FONT_FAMILIES[0], 10), background="#333333",
                  foreground="#cccccc").pack(anchor="w")

        # 运行窗口并等待结果
        preview_window.wait_window()

        # 返回结果
        return getattr(preview_window, 'result', None)

    def save_hotkey_config(self):
        """保存快捷键配置到文件"""
        try:
            config = {
                "hotkey": self.hotkey
            }
            with open(Config.CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showwarning("警告", f"保存快捷键配置失败：{str(e)}")

    def load_hotkey_config(self):
        """从文件加载快捷键配置"""
        try:
            if os.path.exists(Config.CONFIG_FILE):
                with open(Config.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    if "hotkey" in config:
                        self.hotkey = config["hotkey"]
        except Exception as e:
            messagebox.showwarning("警告", f"加载快捷键配置失败：{str(e)}")
            # 使用默认快捷键
            self.hotkey = Config.DEFAULT_HOTKEY

    def set_hotkey(self):
        """设置自定义快捷键"""
        # 检查是否正在捕捉中
        if self.is_capturing:
            messagebox.showinfo("提示", "请先停止当前的捕捉操作，再设置快捷键")
            return

        # 创建自定义对话框
        hotkey_window = tk.Toplevel(self.root)
        hotkey_window.title("设置截图快捷键")
        hotkey_window.geometry("400x180")
        hotkey_window.transient(self.root)
        hotkey_window.grab_set()
        hotkey_window.attributes("-topmost", True)

        # 提示信息
        tip_text = "请输入新的快捷键组合（如：ctrl+alt+s）\n\n注意：\n• 支持的修饰键：ctrl, alt, shift\n• 支持的字母键：a-z\n• 支持的数字键：0-9\n• 支持的功能键：f1-f12\n• 请使用+连接各个键"
        ttk.Label(hotkey_window, text=tip_text, font=(Config.FONT_FAMILIES[0], 10), wraplength=380).pack(pady=15)

        # 输入框
        hotkey_var = tk.StringVar(value=self.hotkey)
        hotkey_entry = ttk.Entry(hotkey_window, textvariable=hotkey_var, width=40, font=(Config.FONT_FAMILIES[0], 12))
        hotkey_entry.pack(pady=10)
        hotkey_entry.focus_set()

        # 确认按钮
        def confirm_hotkey():
            new_hotkey = hotkey_var.get().strip().lower()
            # 简单验证格式
            if not new_hotkey:
                messagebox.showerror("错误", "快捷键不能为空")
                return

            # 检查是否包含无效字符
            valid_modifiers = {'ctrl', 'alt', 'shift'}
            parts = new_hotkey.split('+')
            for part in parts:
                part = part.strip()
                if part and part not in valid_modifiers and not (part.isalpha() and len(part) == 1) and not (
                        part.isdigit() and len(part) == 1) and not (part.startswith('f') and part[1:].isdigit()):
                    messagebox.showerror("错误", "无效的快捷键格式")
                    return

            # 更新快捷键
            self.hotkey = new_hotkey
            self.hotkey_label.config(text=self.hotkey.upper())

            # 保存配置
            self.save_hotkey_config()

            hotkey_window.destroy()
            messagebox.showinfo("成功", f"快捷键已设置为：{self.hotkey.upper()}")

        # 按钮区域
        btn_frame = ttk.Frame(hotkey_window)
        btn_frame.pack(pady=10)
        ttk.Button(btn_frame, text="取消", command=hotkey_window.destroy).pack(side=tk.RIGHT, padx=10)
        ttk.Button(btn_frame, text="确认", command=confirm_hotkey).pack(side=tk.RIGHT, padx=10)

    def show_temp_tip(self, message):
        """显示临时提示窗口"""
        tip_window = tk.Toplevel(self.root)
        tip_window.overrideredirect(True)

        # 计算窗口位置（右下角）
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = len(message) * 10 + 40
        window_height = 50
        x_pos = max(0, screen_width - window_width - 20)
        y_pos = max(0, screen_height - window_height - 40)
        tip_window.geometry(f"{window_width}x{window_height}+{x_pos}+{y_pos}")
        tip_window.configure(bg="#2196f3")
        tip_window.attributes("-topmost", True)

        ttk.Label(
            tip_window,
            text=message,
            background="#2196f3",
            foreground="white",
            padding=10
        ).pack(fill=tk.BOTH, expand=True)

        # 自动关闭
        tip_window.after(Config.TEMP_TIP_DURATION, tip_window.destroy)

    def open_editor_window(self, session):
        """打开会话编辑窗口"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title(f"编辑操作记录 - {session['name']}")
        editor_window.geometry(Config.EDITOR_WINDOW_SIZE)
        editor_window.resizable(True, True)
        editor_window.transient(self.root)

        main_frame = ttk.Frame(editor_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 会话基本信息
        info_frame = ttk.LabelFrame(main_frame, text="操作基本信息", padding="10")
        info_frame.pack(fill=tk.X, pady=5)

        # 名称输入
        ttk.Label(info_frame, text="操作名称：").grid(row=0, column=0, sticky=tk.W, pady=3)
        name_var = tk.StringVar(value=session["name"])
        name_entry = ttk.Entry(info_frame, textvariable=name_var, width=60)
        name_entry.grid(row=0, column=1, sticky=tk.W, pady=3)

        # 描述输入
        ttk.Label(info_frame, text="操作描述：").grid(row=1, column=0, sticky=tk.NW, pady=3)
        desc_text = scrolledtext.ScrolledText(info_frame, width=60, height=4, wrap=tk.WORD)
        desc_text.insert(tk.END, session["description"])
        desc_text.grid(row=1, column=1, sticky=tk.W, pady=3)

        # 只读信息
        info_labels = [
            ("开始时间：", session["start_time"]),
            ("结束时间：", session["end_time"]),
            ("操作时长：", Utils.format_duration(session["duration"])),
            ("截图数量：", str(len(session["captures"])))
        ]
        for i, (label_text, value) in enumerate(info_labels, start=2):
            ttk.Label(info_frame, text=label_text).grid(row=i, column=0, sticky=tk.W, pady=3)
            ttk.Label(info_frame, text=value).grid(row=i, column=1, sticky=tk.W, pady=3)

        # 保存信息按钮
        def save_session_info():
            session["name"] = name_var.get().strip()
            session["description"] = desc_text.get("1.0", tk.END).strip()
            self.save_session(session)
            self.update_history_list()
            messagebox.showinfo("成功", "操作信息已更新")

        ttk.Button(info_frame, text="保存信息", command=save_session_info).grid(row=0, column=2, padx=10, sticky=tk.N)

        # 捕捉记录列表
        ttk.Label(
            main_frame,
            text="截图记录（双击预览，点击操作列编辑/删除）",
            font=(Config.FONT_FAMILIES[0], 12)
        ).pack(anchor=tk.W, pady=5)

        captures_frame = ttk.Frame(main_frame)
        captures_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 截图表格
        columns = ("id", "time", "description", "action")
        captures_tree = ttk.Treeview(captures_frame, columns=columns, show="headings", height=12)
        captures_tree.heading("id", text="序号")
        captures_tree.heading("time", text="截图时间")
        captures_tree.heading("description", text="描述")
        captures_tree.heading("action", text="操作")
        captures_tree.column("id", width=60, anchor=tk.CENTER)
        captures_tree.column("time", width=180)
        captures_tree.column("description", width=400)
        captures_tree.column("action", width=120, anchor=tk.CENTER)
        captures_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 滚动条
        scrollbar = ttk.Scrollbar(captures_frame, orient=tk.VERTICAL, command=captures_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        captures_tree.configure(yscrollcommand=scrollbar.set)

        # 加载截图数据
        for capture in session["captures"]:
            captures_tree.insert("", tk.END, values=(
                capture["id"],
                capture["time"],
                capture["description"],
                "编辑 | 删除"
            ), tags=(capture["id"],))

        # 截图列表操作处理
        def handle_capture_operation(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region != "cell":
                return

            row = captures_tree.identify_row(event.y)
            col = captures_tree.identify_column(event.x)
            if col != "#4":
                return

            capture_id = int(captures_tree.item(row)["tags"][0])
            capture_idx = None
            target_capture = None
            for i, cap in enumerate(session["captures"]):
                if cap["id"] == capture_id:
                    capture_idx = i
                    target_capture = cap
                    break
            if not target_capture:
                return

            x, y, width, height = captures_tree.bbox(row, col)
            if event.x < x + width / 2:
                # 编辑描述
                new_desc = simpledialog.askstring(
                    f"编辑第{capture_id}次截图描述",
                    "请输入新描述：",
                    initialvalue=target_capture["description"]
                )
                if new_desc is not None:
                    target_capture["description"] = new_desc
                    captures_tree.item(row, values=(
                        capture_id,
                        target_capture["time"],
                        new_desc,
                        "编辑 | 删除"
                    ))
                    self.save_session(session)
            else:
                # 删除截图
                if messagebox.askyesno("确认删除", f"确定删除第{capture_id}次截图？\n此操作不可恢复！"):
                    try:
                        if os.path.exists(target_capture["image_path"]):
                            os.remove(target_capture["image_path"])
                    except Exception as e:
                        messagebox.showerror("错误", f"删除图片失败：{str(e)}")
                        return

                    del session["captures"][capture_idx]
                    # 重新编号
                    for i, cap in enumerate(session["captures"]):
                        cap["id"] = i + 1

                    # 刷新表格
                    for item in captures_tree.get_children():
                        captures_tree.delete(item)
                    for cap in session["captures"]:
                        captures_tree.insert("", tk.END, values=(
                            cap["id"],
                            cap["time"],
                            cap["description"],
                            "编辑 | 删除"
                        ), tags=(cap["id"],))
                    self.save_session(session)

        captures_tree.bind("<Button-1>", handle_capture_operation)

        # 双击预览截图
        def preview_screenshot(event):
            region = captures_tree.identify_region(event.x, event.y)
            if region != "cell" or captures_tree.identify_column(event.x) == "#4":
                return

            row = captures_tree.identify_row(event.y)
            capture_id = int(captures_tree.item(row)["tags"][0])
            target_capture = None
            for cap in session["captures"]:
                if cap["id"] == capture_id:
                    target_capture = cap
                    break
            if not target_capture or not os.path.exists(target_capture["image_path"]):
                messagebox.showwarning("提示", "截图文件已损坏或不存在")
                return

            # 使用功能更完整的show_capture_preview替代简单预览
            try:
                # 加载现有图片
                with Image.open(target_capture["image_path"]) as img:
                    # 复制图片以避免修改原始文件
                    screenshot = img.copy()

                # 显示完整预览窗口，并处理返回结果以保存标注
                # 保存原始描述
                original_description = target_capture["description"]

                # 调用预览窗口并获取返回结果，传入原始描述
                result = self.show_capture_preview(screenshot, original_description)

                # 如果用户点击了保存，处理返回的图片和描述
                if result:
                    updated_image, updated_description = result
                    # 更新原始图片
                    try:
                        # 保存更新后的图片到原始路径
                        updated_image.save(target_capture["image_path"])
                        print(f"已保存更新后的图片到: {target_capture['image_path']}")

                        # 如果描述有变化，更新描述
                        if updated_description != original_description:
                            target_capture["description"] = updated_description

                            # 更新界面显示
                            for row_id in captures_tree.get_children():
                                if int(captures_tree.item(row_id)["tags"][0]) == capture_id:
                                    captures_tree.item(row_id, values=(
                                        capture_id,
                                        target_capture["time"],
                                        updated_description,
                                        "编辑 | 删除"
                                    ))

                        # 保存会话数据
                        self.save_session(session)
                        messagebox.showinfo("成功", "标注已保存到截图")
                    except Exception as e:
                        messagebox.showerror("保存错误", f"保存更新后的截图失败：{str(e)}")
                        print(f"保存错误详情: {str(e)}")

            except Exception as e:
                messagebox.showerror("预览错误", f"图片加载失败：{str(e)}")
                return

        captures_tree.bind("<Double-1>", preview_screenshot)

        # 报告生成
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)

        def generate_report():
            if not session["captures"]:
                messagebox.showwarning("提示", "无截图记录，无法生成报告")
                return

            # 选择报告格式
            format_win = tk.Toplevel(editor_window)
            format_win.title("选择报告格式")
            format_win.geometry("350x250")
            format_win.transient(editor_window)
            format_win.grab_set()

            ttk.Label(
                format_win,
                text="请选择报告输出格式：",
                font=(Config.FONT_FAMILIES[0], 12)
            ).pack(pady=20)

            format_var = tk.StringVar(value="docx")

            for fmt in Config.REPORT_FORMATS:
                frame = ttk.Frame(format_win)
                frame.pack(anchor=tk.W, padx=50, pady=3)
                ttk.Radiobutton(
                    frame,
                    text=Config.REPORT_FORMATS[fmt]["title"],
                    variable=format_var,
                    value=fmt
                ).pack(side=tk.LEFT)
                ttk.Label(
                    frame,
                    text=Config.REPORT_FORMATS[fmt]["desc"],
                    font=(Config.FONT_FAMILIES[0], 8),
                    foreground="#666"
                ).pack(side=tk.LEFT, padx=10)

            def confirm_generate():
                report_format = format_var.get()
                format_win.destroy()

                filename = f"{session['name']}_操作报告.{report_format}"
                save_path = filedialog.asksaveasfilename(
                    defaultextension=f".{report_format}",
                    filetypes=[(f"{Config.REPORT_FORMATS[report_format]['title']}", f"*.{report_format}")],
                    initialfile=filename,
                    title=f"保存{report_format.upper()}报告"
                )
                if not save_path:
                    return

                # 显示进度提示
                progress_win = tk.Toplevel(editor_window)
                progress_win.title("生成中")
                progress_win.geometry("300x100")
                progress_win.transient(editor_window)
                progress_win.grab_set()
                ttk.Label(
                    progress_win,
                    text=f"正在生成{report_format.upper()}报告...",
                    font=(Config.FONT_FAMILIES[0], 12)
                ).pack(expand=True)
                progress_win.update()

                try:
                    # 根据格式生成报告
                    if report_format == "docx":
                        DocxReportGenerator.generate(session, save_path)
                    elif report_format == "pdf":
                        PdfReportGenerator.generate(session, save_path)
                    elif report_format == "md":
                        MdReportGenerator.generate(
                            session,
                            save_path,
                            use_relative=True
                        )

                    messagebox.showinfo("成功", f"{report_format.upper()}报告已生成：\n{save_path}")

                    if messagebox.askyesno("打开报告", "是否立即打开生成的报告？"):
                        os.startfile(save_path)
                except Exception as e:
                    messagebox.showerror("生成失败", f"报告生成错误：{str(e)}")
                finally:
                    progress_win.destroy()

            ttk.Button(format_win, text="确认生成", command=confirm_generate).pack(pady=15)

        ttk.Button(btn_frame, text="生成操作报告", command=generate_report).pack(side=tk.LEFT, padx=5)

        # 作废操作
        def discard_operation():
            if messagebox.askyesno(
                    "确认作废",
                    f"确定作废「{session['name']}」操作记录？\n所有截图将被删除，不可恢复！"
            ):
                self.delete_session_by_id(session["id"])
                editor_window.destroy()

        # 首先在截图表格定义后添加多选支持
        captures_tree.configure(selectmode="extended")  # 设置为可多选

        # 然后在按钮区域添加导出按钮（在generate_report函数之后）
        def export_screenshots():
            selected_items = captures_tree.selection()
            if not selected_items:
                messagebox.showwarning("提示", "请先选择要导出的截图")
                return

            # 获取选中的截图
            selected_captures = []
            for item in selected_items:
                capture_id = int(captures_tree.item(item)["tags"][0])
                for cap in session["captures"]:
                    if cap["id"] == capture_id:
                        selected_captures.append(cap)
                        break

            # 选择导出目录
            export_dir = filedialog.askdirectory(title="选择导出目录")
            if not export_dir:
                return

            # 执行导出
            success_count = 0
            fail_count = 0
            fail_list = []

            for cap in selected_captures:
                try:
                    if os.path.exists(cap["image_path"]):
                        # 获取原始文件名
                        filename = os.path.basename(cap["image_path"])
                        # 目标路径
                        target_path = os.path.join(export_dir, filename)

                        # 复制文件
                        with Image.open(cap["image_path"]) as img:
                            img.save(target_path)

                        success_count += 1
                    else:
                        fail_count += 1
                        fail_list.append(f"第{cap['id']}步：文件不存在")
                except Exception as e:
                    fail_count += 1
                    fail_list.append(f"第{cap['id']}步：{str(e)}")

            # 导出结果提示
            result_msg = [f"导出完成：成功 {success_count} 个，失败 {fail_count} 个"]
            if fail_count > 0:
                result_msg.append("\n失败详情：")
                result_msg.extend(fail_list)

            messagebox.showinfo("导出结果", "\n".join(result_msg))

        # 在按钮框架中添加导出按钮（放在生成报告按钮旁边）
        ttk.Button(btn_frame, text="导出选中截图", command=export_screenshots).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="作废此操作", command=discard_operation).pack(side=tk.LEFT, padx=5)

    def _show_image_preview(self, capture, capture_id):
        """显示图片预览窗口，支持滚轮缩放"""
        preview_win = tk.Toplevel(self.root)
        preview_win.title(f"截图预览 - 第{capture_id}次记录")
        preview_win.geometry(Config.PREVIEW_WINDOW_SIZE)

        # 缩放相关变量
        scale_factor = tk.DoubleVar(value=1.0)
        current_image = None
        current_photo = None

        # 滚动条容器
        canvas_frame = ttk.Frame(preview_win)
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        x_scroll = ttk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
        y_scroll = ttk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
        canvas = tk.Canvas(
            canvas_frame,
            xscrollcommand=x_scroll.set,
            yscrollcommand=y_scroll.set,
            bg="#f5f5f5"
        )
        x_scroll.config(command=canvas.xview)
        y_scroll.config(command=canvas.yview)
        x_scroll.pack(side=tk.BOTTOM, fill=tk.X)
        y_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(fill=tk.BOTH, expand=True)

        # 加载图片
        try:
            with Image.open(capture["image_path"]) as img:
                current_image = img.copy()  # 保存原始图像
                self._update_preview_image(canvas, current_image, scale_factor, current_photo)
        except Exception as e:
            messagebox.showerror("预览错误", f"图片加载失败：{str(e)}")
            preview_win.destroy()
            return

        # 缩放函数
        def zoom_image(event):
            nonlocal current_photo
            # 滚轮事件处理
            if event.delta > 0:
                new_scale = min(scale_factor.get() * 1.1, 5.0)  # 最大放大5倍
            else:
                new_scale = max(scale_factor.get() / 1.1, 0.1)  # 最小缩小到0.1倍

            scale_factor.set(new_scale)
            current_photo = self._update_preview_image(canvas, current_image, scale_factor, current_photo)

        # 绑定滚轮事件
        canvas.bind("<MouseWheel>", zoom_image)  # Windows
        canvas.bind("<Button-4>", zoom_image)  # Linux
        canvas.bind("<Button-5>", zoom_image)  # Linux

        ttk.Label(
            preview_win,
            text=f"描述：{capture['description']}",
            font=(Config.FONT_FAMILIES[0], 10)
        ).pack(pady=5)

    def _update_preview_image(self, canvas, original_img, scale_factor, old_photo):
        """更新预览图片（缩放处理）"""
        scale = scale_factor.get()
        img_width, img_height = original_img.size
        new_width = int(img_width * scale)
        new_height = int(img_height * scale)

        # 缩放图片（保持高质量）
        resized_img = original_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        photo = ImageTk.PhotoImage(resized_img)

        # 清除旧图像并显示新图像
        canvas.delete("all")
        canvas.create_image(0, 0, anchor=tk.NW, image=photo)
        canvas.config(scrollregion=(0, 0, new_width, new_height))

        # 保存引用防止垃圾回收
        canvas.image = photo
        return photo

    def save_current_session(self):
        """保存当前会话"""
        if self.current_session_id and self.current_session:
            self.save_session(self.current_session)

    def save_session(self, session):
        """保存会话到文件"""
        session_path = os.path.join(self.sessions_dir, f"{session['id']}.json")
        try:
            with open(session_path, "w", encoding="utf-8") as f:
                json.dump(session, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("保存错误", f"会话保存失败：{str(e)}")

    def load_history_sessions(self):
        """加载历史会话"""
        self.history_sessions = []
        try:
            for filename in os.listdir(self.sessions_dir):
                if filename.endswith(".json"):
                    session_id = filename[:-5]
                    session_path = os.path.join(self.sessions_dir, filename)
                    with open(session_path, "r", encoding="utf-8") as f:
                        session = json.load(f)
                        self.history_sessions.append(session)
            # 按开始时间排序（最新的在前）
            self.history_sessions.sort(key=lambda x: x["start_time"], reverse=True)
        except Exception as e:
            messagebox.showerror("加载错误", f"历史记录加载失败：{str(e)}")

    def update_history_list(self):
        """更新历史记录列表"""
        # 清空现有项
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)

        # 添加所有会话
        for session in self.history_sessions:
            self.history_tree.insert("", tk.END, values=(
                session["name"],
                session["start_time"],
                Utils.format_duration(session["duration"]),
                len(session["captures"])
            ), tags=(session["id"],))

    def open_session(self, event):
        """打开选中的会话"""
        selected_item = self.history_tree.selection()
        if not selected_item:
            return

        session_id = self.history_tree.item(selected_item[0])["tags"][0]
        for session in self.history_sessions:
            if session["id"] == session_id:
                self.open_editor_window(session)
                break

    def on_main_window_close(self):
        """主窗口关闭处理"""
        if self.is_capturing:
            if messagebox.askyesno("提示", "当前正在捕捉中，确定要退出吗？"):
                self.stop_capture()
                self.root.destroy()
        else:
            self.root.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = ScreenCaptureTool(root)
    root.mainloop()